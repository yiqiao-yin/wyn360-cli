"""Document readers with intelligent chunking, summarization, and retrieval.

This module provides intelligent document reading capabilities for Excel, Word, and PDF files.
Instead of simple truncation, it implements a chunking + summarization + tagging + retrieval
system that allows handling arbitrarily large documents.

Architecture:
    Document → Chunk → Summarize → Tag → Cache → Query → Retrieve

Phase 1: Core Infrastructure (v0.3.26)
"""

import hashlib
import json
import time
import re
import base64
import io
from pathlib import Path
from typing import Optional, List, Dict, Tuple, Any, Union
from dataclasses import dataclass, asdict

# Optional dependencies (graceful fallback)
try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False
    openpyxl = None

try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    Document = None

try:
    import pymupdf  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    pymupdf = None

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False
    pdfplumber = None

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
    Image = None

# Import Anthropic at module level for easier mocking in tests
try:
    from anthropic import Anthropic
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False
    Anthropic = None

# Import numpy for embeddings (Phase 5.2.1)
try:
    import numpy as np
    HAS_NUMPY = True
except ImportError:
    HAS_NUMPY = False
    np = None


# ============================================================================
# Token Counting Utilities
# ============================================================================

def count_tokens(text: str) -> int:
    """
    Estimate token count for text.

    Uses rough approximation: 1 token ≈ 4 characters.
    Same as browser_use.py for consistency.

    Args:
        text: Text to count tokens for

    Returns:
        Estimated token count
    """
    return len(text) // 4


# ============================================================================
# Data Classes
# ============================================================================

@dataclass
class ChunkMetadata:
    """Metadata for a document chunk."""
    chunk_id: str
    position: Dict[str, Any]  # start, end positions (varies by doc type)
    summary: str
    tags: List[str]
    token_count: int
    summary_tokens: int
    tag_tokens: int
    # Optional type-specific fields
    sheet_name: Optional[str] = None  # Excel
    section_title: Optional[str] = None  # Word
    page_range: Optional[Tuple[int, int]] = None  # PDF
    embedding: Optional[List[float]] = None  # Phase 5.2: Semantic embeddings


@dataclass
class DocumentMetadata:
    """Metadata for a cached document."""
    file_path: str
    file_hash: str
    file_size: int
    total_tokens: int
    chunk_count: int
    chunk_size: int
    created_at: float
    ttl: int
    doc_type: str  # "excel" | "word" | "pdf"


# ============================================================================
# ImageProcessor Class (Phase 5.1)
# ============================================================================

class ImageProcessor:
    """
    Process images from documents using Claude Vision API.

    Features:
    - Extract images from Word and PDF documents
    - Send images to Claude Vision API for description
    - Support batch processing for efficiency
    - Track vision API costs separately
    - Detect image types (chart, diagram, photo, etc.)

    Usage:
        processor = ImageProcessor(api_key="your_key")
        description = await processor.describe_image(image_data, context={...})
    """

    def __init__(self, api_key: str, model: str = "claude-3-5-sonnet-20241022"):
        """
        Initialize image processor.

        Args:
            api_key: Anthropic API key
            model: Claude model to use (must support vision)
        """
        if not HAS_ANTHROPIC:
            raise ImportError("anthropic library required for vision mode")

        self.api_key = api_key
        self.model = model
        self.client = Anthropic(api_key=api_key)

    async def describe_image(
        self,
        image_data: bytes,
        context: Optional[Dict[str, Any]] = None,
        max_tokens: int = 300
    ) -> Dict[str, Any]:
        """
        Describe an image using Claude Vision API.

        Args:
            image_data: Image bytes (PNG, JPEG, GIF, WebP)
            context: Optional context about the image (page number, document type, etc.)
            max_tokens: Maximum tokens for description

        Returns:
            {
                "description": "Image description...",
                "image_type": "chart" | "diagram" | "photo" | "screenshot" | "other",
                "tokens_used": 150,
                "confidence": "high" | "medium" | "low"
            }
        """
        try:
            # Detect image format
            image_format = self._detect_image_format(image_data)
            if not image_format:
                return {
                    "description": "[Image: format not supported]",
                    "image_type": "other",
                    "tokens_used": 0,
                    "confidence": "low"
                }

            # Encode image to base64
            image_base64 = base64.b64encode(image_data).decode('utf-8')

            # Build prompt based on context
            prompt = self._build_image_prompt(context)

            # Call Claude Vision API
            message = self.client.messages.create(
                model=self.model,
                max_tokens=max_tokens,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": f"image/{image_format}",
                                    "data": image_base64
                                }
                            },
                            {
                                "type": "text",
                                "text": prompt
                            }
                        ]
                    }
                ]
            )

            # Extract description
            description = message.content[0].text if message.content else "[No description available]"

            # Detect image type from description
            image_type = self._detect_image_type(description)

            # Calculate tokens used
            tokens_used = message.usage.input_tokens + message.usage.output_tokens

            return {
                "description": description,
                "image_type": image_type,
                "tokens_used": tokens_used,
                "confidence": "high"
            }

        except Exception as e:
            return {
                "description": f"[Image: error processing - {str(e)}]",
                "image_type": "other",
                "tokens_used": 0,
                "confidence": "low"
            }

    async def describe_images_batch(
        self,
        images: List[Dict[str, Any]],
        context: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """
        Describe multiple images efficiently.

        Args:
            images: List of {"data": bytes, "context": dict} dicts
            context: Global context for all images

        Returns:
            List of description results
        """
        results = []

        for img in images:
            img_context = {**(context or {}), **(img.get("context", {}))}
            result = await self.describe_image(
                image_data=img["data"],
                context=img_context
            )
            results.append(result)

        return results

    def _detect_image_format(self, image_data: bytes) -> Optional[str]:
        """Detect image format from bytes."""
        # Check magic bytes
        if image_data.startswith(b'\x89PNG'):
            return "png"
        elif image_data.startswith(b'\xff\xd8\xff'):
            return "jpeg"
        elif image_data.startswith(b'GIF'):
            return "gif"
        elif image_data.startswith(b'RIFF') and b'WEBP' in image_data[:20]:
            return "webp"
        else:
            # Try using PIL if available
            if HAS_PIL:
                try:
                    img = Image.open(io.BytesIO(image_data))
                    fmt = img.format.lower()
                    if fmt in ['png', 'jpeg', 'jpg', 'gif', 'webp']:
                        return 'jpeg' if fmt == 'jpg' else fmt
                except:
                    pass
            return None

    def _build_image_prompt(self, context: Optional[Dict[str, Any]] = None) -> str:
        """Build prompt for image description based on context."""
        base_prompt = "Describe this image concisely. "

        if not context:
            return base_prompt + "Focus on the main content and purpose of the image."

        doc_type = context.get("doc_type", "")
        page_info = context.get("page_number", "")
        section_info = context.get("section_title", "")

        # Customize prompt based on document type
        if doc_type == "pdf" or doc_type == "word":
            base_prompt += "This image is from a document. "

        if page_info:
            base_prompt += f"(Page {page_info}) "

        if section_info:
            base_prompt += f"(Section: {section_info}) "

        base_prompt += "Identify if this is a chart, diagram, photo, screenshot, or other type of image. "
        base_prompt += "Describe its content and purpose clearly and concisely."

        return base_prompt

    def _detect_image_type(self, description: str) -> str:
        """Detect image type from description."""
        description_lower = description.lower()

        if any(word in description_lower for word in ['chart', 'graph', 'plot', 'bar', 'line', 'pie']):
            return "chart"
        elif any(word in description_lower for word in ['diagram', 'flowchart', 'schematic', 'blueprint']):
            return "diagram"
        elif any(word in description_lower for word in ['screenshot', 'screen capture', 'interface']):
            return "screenshot"
        elif any(word in description_lower for word in ['photo', 'photograph', 'picture', 'image of']):
            return "photo"
        else:
            return "other"

    def format_image_markdown(self, description: str, image_type: str, image_index: int = None) -> str:
        """Format image description as markdown."""
        type_emoji = {
            "chart": "📊",
            "diagram": "📐",
            "screenshot": "🖥️",
            "photo": "📷",
            "other": "🖼️"
        }

        emoji = type_emoji.get(image_type, "🖼️")
        prefix = f"Image {image_index}" if image_index else "Image"

        return f"{emoji} **[{prefix}]:** {description}"


# ============================================================================
# OCRProcessor Class (Phase 5.3.1)
# ============================================================================

# Try importing OCR dependencies
HAS_PYTESSERACT = True
HAS_PDF2IMAGE = True
try:
    import pytesseract
    from PIL import Image, ImageEnhance
except ImportError:
    HAS_PYTESSERACT = False

try:
    from pdf2image import convert_from_path
except ImportError:
    HAS_PDF2IMAGE = False


class OCRProcessor:
    """
    Extract text from images using Tesseract OCR.

    Features:
    - Detect scanned PDF pages (image-based)
    - Extract text from images with confidence scoring
    - Image preprocessing (grayscale, contrast enhancement)
    - Language detection and multi-language support
    - Quality assessment

    Usage:
        processor = OCRProcessor(language="eng")
        result = processor.extract_text(image)
        # {"text": "...", "confidence": 85.5, "word_count": 120}
    """

    def __init__(self, language: str = "eng"):
        """
        Initialize OCR processor.

        Args:
            language: Tesseract language code (eng, spa, fra, deu, etc.)
        """
        self.language = language
        self._check_tesseract()

    def _check_tesseract(self):
        """Verify Tesseract is installed."""
        if not HAS_PYTESSERACT:
            raise RuntimeError(
                "Tesseract OCR not installed. "
                "Install: pip install pytesseract Pillow\n"
                "System: apt-get install tesseract-ocr (Linux) "
                "or brew install tesseract (Mac)"
            )

        try:
            # Test if tesseract binary is accessible
            pytesseract.get_tesseract_version()
        except Exception as e:
            raise RuntimeError(
                f"Tesseract binary not found: {e}\n"
                "Install: apt-get install tesseract-ocr (Linux) "
                "or brew install tesseract (Mac)"
            )

    def extract_text(
        self,
        image: "Image.Image",
        preprocess: bool = True
    ) -> Dict[str, Any]:
        """
        Extract text from image using OCR.

        Args:
            image: PIL Image object
            preprocess: Apply preprocessing (grayscale, contrast)

        Returns:
            {
                "text": str,
                "confidence": float (0-100),
                "word_count": int
            }
        """
        if preprocess:
            image = self._preprocess_image(image)

        # Extract text with confidence data
        data = pytesseract.image_to_data(
            image,
            lang=self.language,
            output_type=pytesseract.Output.DICT
        )

        # Calculate average confidence (excluding -1 values)
        confidences = [c for c in data['conf'] if c != -1]
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0

        # Extract text
        text = pytesseract.image_to_string(image, lang=self.language)

        return {
            "text": text.strip(),
            "confidence": avg_confidence,
            "word_count": len(text.split())
        }

    def _preprocess_image(self, image: "Image.Image") -> "Image.Image":
        """
        Preprocess image for better OCR results.

        Args:
            image: PIL Image object

        Returns:
            Preprocessed PIL Image
        """
        from PIL import Image, ImageEnhance

        # Convert to grayscale
        image = image.convert('L')

        # Increase contrast
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(2.0)

        return image

    def is_scanned_page(
        self,
        page_image: "Image.Image",
        text_from_pdf: str,
        min_words_threshold: int = 20
    ) -> bool:
        """
        Detect if PDF page is scanned (image-based).

        Args:
            page_image: Page rendered as image
            text_from_pdf: Text extracted directly from PDF
            min_words_threshold: Minimum words to consider page as scanned

        Returns:
            True if page appears to be scanned
        """
        # If PDF has very little text but page has visual content, likely scanned
        if len(text_from_pdf.strip()) < 10:
            # Try quick OCR to see if there's text
            result = self.extract_text(page_image, preprocess=False)
            return result["word_count"] > min_words_threshold

        return False

    def assess_quality(self, ocr_result: Dict[str, Any]) -> str:
        """
        Assess OCR quality based on confidence score.

        Args:
            ocr_result: Result from extract_text()

        Returns:
            Quality rating: "excellent", "good", "fair", "poor"
        """
        confidence = ocr_result["confidence"]

        if confidence > 90:
            return "excellent"
        elif confidence > 75:
            return "good"
        elif confidence > 60:
            return "fair"
        else:
            return "poor"


# ============================================================================
# EmbeddingModel Class (Phase 5.2.1)
# ============================================================================

class EmbeddingModel:
    """
    Wrapper for embedding models with support for local and API-based embeddings.

    Supported providers:
    - "local": sentence-transformers models (default: all-MiniLM-L6-v2)
    - "claude": Claude embeddings via Anthropic API (voyage-3 model)

    Local model (default):
    - all-MiniLM-L6-v2 (22MB, 384 dimensions)
    - Fast inference (~0.01s per sentence on CPU)
    - No API costs

    Claude embeddings:
    - High quality (1024 dimensions)
    - API costs apply: $0.0001 per 1K tokens
    - Requires Anthropic API key
    """

    # Whitelist of safe sentence-transformers models
    SAFE_LOCAL_MODELS = {
        "all-MiniLM-L6-v2",  # Default, lightweight
        "all-mpnet-base-v2",  # Higher quality, larger
        "paraphrase-MiniLM-L6-v2",  # Good for paraphrase detection
        "multi-qa-MiniLM-L6-cos-v1",  # Optimized for Q&A
    }

    def __init__(
        self,
        provider: str = "local",
        model_name: str = "all-MiniLM-L6-v2",
        api_key: Optional[str] = None
    ):
        """
        Initialize embedding model.

        Args:
            provider: "local" for sentence-transformers, "claude" for Anthropic API
            model_name: Model name (only used for "local" provider)
            api_key: Anthropic API key (only needed for "claude" provider)
        """
        self.provider = provider
        self.model_name = model_name
        self.api_key = api_key
        self.model = None
        self._initialized = False

        # Validate provider
        if self.provider not in ["local", "claude"]:
            raise ValueError(
                f"Invalid provider '{self.provider}'. "
                "Must be 'local' or 'claude'."
            )

        # Validate local model name (security: prevent arbitrary model loading)
        if self.provider == "local" and self.model_name not in self.SAFE_LOCAL_MODELS:
            raise ValueError(
                f"Model '{self.model_name}' not in safe model list. "
                f"Allowed models: {', '.join(self.SAFE_LOCAL_MODELS)}"
            )

        # Validate API key for Claude
        if self.provider == "claude" and not self.api_key:
            raise ValueError(
                "api_key required for 'claude' provider"
            )

    def _lazy_load(self):
        """Lazy load model on first use to avoid startup overhead."""
        if not self._initialized:
            if self.provider == "local":
                try:
                    from sentence_transformers import SentenceTransformer
                    self.model = SentenceTransformer(self.model_name)
                    self._initialized = True
                except ImportError as e:
                    raise ImportError(
                        "sentence-transformers not installed. "
                        "Install with: pip install sentence-transformers torch numpy"
                    ) from e
            elif self.provider == "claude":
                # No lazy load needed for API-based embeddings
                self._initialized = True

    def encode(self, texts: Union[str, List[str]]) -> np.ndarray:
        """
        Encode text(s) into embeddings.

        Args:
            texts: Single text or list of texts

        Returns:
            Embeddings as numpy array (n_texts, embedding_dim)
        """
        self._lazy_load()

        if isinstance(texts, str):
            texts = [texts]

        if self.provider == "local":
            embeddings = self.model.encode(texts, convert_to_numpy=True)
            return embeddings

        elif self.provider == "claude":
            # Use Claude embeddings API
            import numpy as np
            from anthropic import Anthropic

            client = Anthropic(api_key=self.api_key)

            # Note: As of now, Anthropic doesn't have a dedicated embeddings API
            # This is a placeholder for when they add it
            # For now, we'll raise an error with a helpful message
            raise NotImplementedError(
                "Claude embeddings API not yet available. "
                "Use provider='local' with sentence-transformers models instead."
            )

            # Future implementation when Anthropic adds embeddings:
            # response = client.embeddings.create(
            #     model="claude-embeddings-v1",
            #     texts=texts
            # )
            # embeddings = np.array([emb.embedding for emb in response.data])
            # return embeddings

    def compute_similarity(
        self,
        query_embedding: np.ndarray,
        chunk_embeddings: np.ndarray
    ) -> np.ndarray:
        """
        Compute cosine similarity between query and chunks.

        Args:
            query_embedding: Query embedding (1, embedding_dim) or (embedding_dim,)
            chunk_embeddings: Chunk embeddings (n_chunks, embedding_dim)

        Returns:
            Similarity scores (n_chunks,)
        """
        import numpy as np

        # Ensure query_embedding is 1D
        if query_embedding.ndim > 1:
            query_embedding = query_embedding.flatten()

        # Normalize embeddings
        query_norm = query_embedding / np.linalg.norm(query_embedding)
        chunk_norms = chunk_embeddings / np.linalg.norm(
            chunk_embeddings, axis=1, keepdims=True
        )

        # Cosine similarity
        similarities = np.dot(chunk_norms, query_norm).flatten()
        return similarities


# ============================================================================
# DocumentChunker Class
# ============================================================================

class DocumentChunker:
    """
    Intelligently chunk documents into manageable pieces.

    Phase 5.7: Enhanced with adaptive sizing, overlap, content-awareness, and quality scoring.
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        adaptive_sizing: bool = False,
        overlap_tokens: int = 0,
        content_aware: bool = False,
        quality_threshold: float = 0.0
    ):
        """
        Initialize chunker.

        Args:
            chunk_size: Target tokens per chunk (default: 1000)
            adaptive_sizing: Enable adaptive chunk sizes based on content density (Phase 5.7.1)
            overlap_tokens: Number of tokens to overlap between chunks (Phase 5.7.2)
            content_aware: Don't split tables, code blocks, lists (Phase 5.7.3)
            quality_threshold: Minimum quality score for chunks, 0-1 (Phase 5.7.4)
        """
        self.chunk_size = chunk_size
        self.adaptive_sizing = adaptive_sizing
        self.overlap_tokens = overlap_tokens
        self.content_aware = content_aware
        self.quality_threshold = quality_threshold

    def _calculate_adaptive_size(self, text: str) -> int:
        """
        Calculate adaptive chunk size based on content density.

        Phase 5.7.1: Adaptive Sizing

        Dense content (tables, lists, code) gets smaller chunks (~500 tokens).
        Sparse content (prose) gets larger chunks (~1500 tokens).

        Args:
            text: Text to analyze

        Returns:
            Recommended chunk size in tokens
        """
        if not text:
            return self.chunk_size

        # Detect tables (markdown or simple text tables)
        table_pattern = r'\|.*\|.*\|'
        has_tables = bool(re.search(table_pattern, text))

        # Detect lists (markdown style)
        list_pattern = r'^\s*[-*•]\s|\d+\.\s'
        has_lists = bool(re.search(list_pattern, text, re.MULTILINE))

        # Detect code blocks
        code_pattern = r'```|^\s{4,}'
        has_code = bool(re.search(code_pattern, text, re.MULTILINE))

        # Dense content: smaller chunks for better granularity
        if has_tables or has_lists or has_code:
            return max(500, self.chunk_size // 2)

        # Sparse content: larger chunks for better context
        return min(1500, int(self.chunk_size * 1.5))

    def chunk_by_tokens(self, text: str, preserve_boundaries: bool = True) -> List[str]:
        """
        Chunk text by token count.

        Phase 5.7.1: Now supports adaptive sizing based on content density.
        Phase 5.7.3: Now supports content-aware boundaries (don't split tables, code, lists).
        Phase 5.7.4: Now supports quality scoring and filtering.

        Args:
            text: Text to chunk
            preserve_boundaries: Try to break at paragraph boundaries

        Returns:
            List of text chunks (filtered by quality if quality_threshold > 0)
        """
        # Use adaptive size if enabled, otherwise use default chunk_size
        if self.adaptive_sizing:
            chunk_size = self._calculate_adaptive_size(text)
        else:
            chunk_size = self.chunk_size

        target_chars = chunk_size * 4  # ~4 chars per token

        # Create chunks based on selected strategy
        if self.content_aware:
            chunks = self._chunk_respecting_blocks(text, target_chars)
        elif preserve_boundaries:
            chunks = self._chunk_with_boundaries(text, target_chars)
        else:
            chunks = self._chunk_simple(text, target_chars)

        # Apply quality filtering if threshold is set
        if self.quality_threshold > 0.0:
            chunks = self._apply_quality_filtering(chunks)

        return chunks

    def _chunk_simple(self, text: str, target_chars: int) -> List[str]:
        """
        Simple chunking by character count.

        Phase 5.7.2: Now supports overlapping chunks when overlap_tokens > 0.
        """
        if not text:
            return [""]

        # If overlap is enabled, use overlapping chunking
        if self.overlap_tokens > 0:
            return self._chunk_with_overlap(text, target_chars)

        # Original non-overlapping chunking
        chunks = []
        for i in range(0, len(text), target_chars):
            chunks.append(text[i:i + target_chars])
        return chunks

    def _chunk_with_overlap(self, text: str, target_chars: int) -> List[str]:
        """
        Chunk text with overlapping regions.

        Phase 5.7.2: Overlapping Chunks

        Creates chunks with specified overlap to preserve context across boundaries.
        For example, with overlap_tokens=200:
        - chunk1: 0-1000
        - chunk2: 800-1800 (200 token overlap)
        - chunk3: 1600-2600

        Args:
            text: Text to chunk
            target_chars: Target characters per chunk (~4 chars per token)

        Returns:
            List of overlapping text chunks
        """
        if not text:
            return [""]

        chunks = []
        overlap_chars = self.overlap_tokens * 4  # ~4 chars per token

        # Ensure overlap doesn't exceed chunk size
        overlap_chars = min(overlap_chars, target_chars // 2)

        position = 0
        while position < len(text):
            # Extract chunk
            end_position = min(position + target_chars, len(text))
            chunk = text[position:end_position]
            chunks.append(chunk)

            # Move position forward, accounting for overlap
            # Next chunk starts at (current_end - overlap)
            position = end_position - overlap_chars

            # Prevent infinite loop if we're at the end
            if end_position == len(text):
                break

            # Ensure we're making progress
            if position >= end_position:
                position = end_position

        return chunks

    def _chunk_with_boundaries(self, text: str, target_chars: int) -> List[str]:
        """
        Chunk with respect to paragraph boundaries.

        Phase 5.7.2: Now supports overlapping chunks when overlap_tokens > 0.
        """
        # If overlap is enabled, first create overlapping chunks, then try to adjust boundaries
        if self.overlap_tokens > 0:
            # For boundary-aware overlapping, we use a hybrid approach:
            # Create overlapping chunks, but try to break at paragraph boundaries
            return self._chunk_with_overlap(text, target_chars)

        # Original non-overlapping boundary-aware chunking
        chunks = []
        paragraphs = text.split('\n\n')

        current_chunk = ""
        for para in paragraphs:
            if len(current_chunk) + len(para) <= target_chars:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = para + "\n\n"

        if current_chunk:
            chunks.append(current_chunk.strip())

        return chunks

    def _detect_content_blocks(self, text: str) -> List[Dict[str, Any]]:
        """
        Detect content blocks (tables, code blocks, lists) that shouldn't be split.

        Phase 5.7.3: Content-Aware Boundaries

        Returns:
            List of block dicts with {type, start, end, content}
        """
        blocks = []

        # Detect code blocks (triple backticks)
        code_pattern = r'```[\s\S]*?```'
        for match in re.finditer(code_pattern, text):
            blocks.append({
                'type': 'code',
                'start': match.start(),
                'end': match.end(),
                'content': match.group()
            })

        # Detect tables (markdown style)
        # Look for lines with multiple | characters
        lines = text.split('\n')
        in_table = False
        table_start = 0
        table_lines = []

        for i, line in enumerate(lines):
            is_table_line = line.count('|') >= 2
            if is_table_line:
                if not in_table:
                    in_table = True
                    table_start = sum(len(l) + 1 for l in lines[:i])  # +1 for \n
                    table_lines = [line]
                else:
                    table_lines.append(line)
            else:
                if in_table:
                    # Table ended
                    table_content = '\n'.join(table_lines)
                    table_end = table_start + len(table_content)
                    blocks.append({
                        'type': 'table',
                        'start': table_start,
                        'end': table_end,
                        'content': table_content
                    })
                    in_table = False
                    table_lines = []

        # Handle table at end of text
        if in_table:
            table_content = '\n'.join(table_lines)
            table_end = table_start + len(table_content)
            blocks.append({
                'type': 'table',
                'start': table_start,
                'end': table_end,
                'content': table_content
            })

        # Detect lists (consecutive list items)
        list_pattern = r'(?:^\s*[-*•]\s.+$|^\s*\d+\.\s.+$)+'
        for match in re.finditer(list_pattern, text, re.MULTILINE):
            # Only add if it's at least 2 lines (actual list)
            if match.group().count('\n') >= 1:
                blocks.append({
                    'type': 'list',
                    'start': match.start(),
                    'end': match.end(),
                    'content': match.group()
                })

        # Sort blocks by start position
        blocks.sort(key=lambda b: b['start'])

        return blocks

    def _chunk_respecting_blocks(self, text: str, target_chars: int) -> List[str]:
        """
        Chunk text while respecting content blocks (don't split tables, code, lists).

        Phase 5.7.3: Content-Aware Boundaries

        Args:
            text: Text to chunk
            target_chars: Target characters per chunk

        Returns:
            List of chunks that don't split content blocks
        """
        if not text:
            return [""]

        # Detect content blocks
        blocks = self._detect_content_blocks(text)

        if not blocks:
            # No blocks detected, use regular chunking
            if self.overlap_tokens > 0:
                return self._chunk_with_overlap(text, target_chars)
            else:
                return self._chunk_simple(text, target_chars)

        chunks = []
        current_chunk = ""
        position = 0

        for block in blocks:
            # Add text before this block to current chunk
            text_before = text[position:block['start']]

            # Check if adding this text would exceed target
            if text_before:
                if len(current_chunk) + len(text_before) <= target_chars:
                    current_chunk += text_before
                else:
                    # Split the text_before if needed
                    if current_chunk:
                        chunks.append(current_chunk)
                    # Chunk the text_before part
                    temp_chunks = self._chunk_simple(text_before, target_chars)
                    chunks.extend(temp_chunks[:-1])  # Add all but last
                    current_chunk = temp_chunks[-1] if temp_chunks else ""

            # Now handle the block itself
            block_content = block['content']

            # If block is too large, put it in its own chunk(s)
            if len(block_content) > target_chars:
                if current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = ""
                # Split large block into multiple chunks
                # (We tried not to split it, but it's too big)
                block_chunks = self._chunk_simple(block_content, target_chars)
                chunks.extend(block_chunks)
            # If adding block would exceed target, start new chunk
            elif len(current_chunk) + len(block_content) > target_chars:
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = block_content
            else:
                # Add block to current chunk
                current_chunk += block_content

            position = block['end']

        # Add remaining text after last block
        if position < len(text):
            remaining_text = text[position:]
            if len(current_chunk) + len(remaining_text) <= target_chars:
                current_chunk += remaining_text
            else:
                if current_chunk:
                    chunks.append(current_chunk)
                # Chunk the remaining text
                remaining_chunks = self._chunk_simple(remaining_text, target_chars)
                chunks.extend(remaining_chunks)
                current_chunk = ""

        # Add final chunk if any
        if current_chunk:
            chunks.append(current_chunk)

        return chunks if chunks else [""]

    def score_chunk_quality(self, chunk: str) -> float:
        """
        Score chunk quality based on coherence, completeness, and independence.

        Phase 5.7.4: Quality Scoring

        Scores range from 0.0 (low quality) to 1.0 (high quality).

        Quality criteria:
        - Coherence: Does the chunk form a complete thought?
        - Completeness: Are relevant details present?
        - Independence: Can the chunk be understood alone?

        Args:
            chunk: Text chunk to score

        Returns:
            Quality score (0.0 to 1.0)
        """
        if not chunk or len(chunk.strip()) == 0:
            return 0.0

        score = 1.0
        penalties = []

        # Coherence checks
        # 1. Check for incomplete sentences (starts/ends mid-sentence)
        chunk_stripped = chunk.strip()

        # Does it start with lowercase (likely mid-sentence)?
        if chunk_stripped and chunk_stripped[0].islower():
            penalties.append(('starts_lowercase', 0.1))

        # Does it end without proper punctuation?
        if chunk_stripped and chunk_stripped[-1] not in '.!?"\'"':
            penalties.append(('no_end_punctuation', 0.1))

        # Completeness checks
        # 2. Check for orphaned list items (list with only 1 item)
        list_item_pattern = r'^\s*[-*•]\s|\d+\.\s'
        list_matches = re.findall(list_item_pattern, chunk, re.MULTILINE)
        if len(list_matches) == 1:
            penalties.append(('orphaned_list_item', 0.15))

        # 3. Check for incomplete tables (header without data, or vice versa)
        lines = chunk.split('\n')
        table_lines = [line for line in lines if line.count('|') >= 2]
        if len(table_lines) == 1:
            # Only one table row (likely incomplete)
            penalties.append(('incomplete_table', 0.15))

        # 4. Check for incomplete code blocks
        code_block_starts = chunk.count('```')
        if code_block_starts % 2 != 0:
            # Odd number of backticks = incomplete code block
            penalties.append(('incomplete_code_block', 0.2))

        # Independence checks
        # 5. Check for unresolved references (e.g., "this", "that", "it" at start)
        first_words = chunk_stripped.split()[:3] if chunk_stripped else []
        pronoun_references = ['this', 'that', 'it', 'these', 'those', 'they']
        if any(word.lower() in pronoun_references for word in first_words):
            penalties.append(('starts_with_reference', 0.1))

        # 6. Check for very short chunks (< 50 chars, likely incomplete)
        if len(chunk_stripped) < 50:
            penalties.append(('too_short', 0.15))

        # 7. Check for extremely long chunks (> 20,000 chars, likely not properly chunked)
        if len(chunk_stripped) > 20000:
            penalties.append(('too_long', 0.1))

        # Apply penalties (cap at 0.0)
        total_penalty = sum(penalty for _, penalty in penalties)
        score = max(0.0, score - total_penalty)

        return score

    def _apply_quality_filtering(self, chunks: List[str]) -> List[str]:
        """
        Filter chunks based on quality threshold.

        Phase 5.7.4: Quality Scoring

        Args:
            chunks: List of text chunks

        Returns:
            Filtered list of chunks with quality >= threshold
        """
        if self.quality_threshold <= 0.0:
            # No filtering if threshold is 0
            return chunks

        filtered_chunks = []
        for chunk in chunks:
            quality = self.score_chunk_quality(chunk)
            if quality >= self.quality_threshold:
                filtered_chunks.append(chunk)

        # If all chunks filtered out, return at least one (the best quality chunk)
        if not filtered_chunks and chunks:
            # Find chunk with highest quality
            best_chunk = max(chunks, key=lambda c: self.score_chunk_quality(c))
            filtered_chunks = [best_chunk]

        return filtered_chunks if filtered_chunks else chunks

    def chunk_by_structure(
        self,
        sections: List[Dict[str, str]],
        target_chunk_tokens: int = 1000
    ) -> List[Dict[str, Any]]:
        """
        Chunk by document structure (headings/sections).

        Used for Word documents where we want to keep sections together.

        Args:
            sections: List of {title, content, level} dicts
            target_chunk_tokens: Target tokens per chunk

        Returns:
            List of chunk dicts with aggregated sections
        """
        chunks = []
        current_chunk = {
            "title": "",
            "content": "",
            "sections": []
        }
        current_tokens = 0

        for section in sections:
            section_text = f"# {section['title']}\n\n{section['content']}\n\n"
            section_tokens = count_tokens(section_text)

            if current_tokens + section_tokens <= target_chunk_tokens:
                current_chunk["content"] += section_text
                current_chunk["sections"].append(section["title"])
                current_tokens += section_tokens
            else:
                if current_chunk["content"]:
                    chunks.append(current_chunk.copy())

                current_chunk = {
                    "title": section["title"],
                    "content": section_text,
                    "sections": [section["title"]]
                }
                current_tokens = section_tokens

        if current_chunk["content"]:
            chunks.append(current_chunk)

        return chunks


# ============================================================================
# ChunkSummarizer Class
# ============================================================================

class ChunkSummarizer:
    """Summarize chunks and generate tags using Claude API.

    Phase 5.2.2: Now includes semantic embedding generation for improved retrieval.
    """

    def __init__(
        self,
        api_key: str,
        model: str = "claude-3-5-haiku-20241022",
        enable_embeddings: bool = True,
        embedding_provider: str = "local",
        embedding_model: str = "all-MiniLM-L6-v2"
    ):
        """
        Initialize summarizer.

        Args:
            api_key: Anthropic API key
            model: Claude model to use (default: Haiku for cost efficiency)
            enable_embeddings: Enable semantic embeddings (Phase 5.2.2)
            embedding_provider: "local" or "claude" (default: local)
            embedding_model: Model name for local embeddings
        """
        self.api_key = api_key
        self.model = model
        self.summary_tokens = 100
        self.tag_count = 8
        self.enable_embeddings = enable_embeddings
        self.embedding_model = None

        # Initialize embedding model if enabled
        if self.enable_embeddings:
            try:
                self.embedding_model = EmbeddingModel(
                    provider=embedding_provider,
                    model_name=embedding_model,
                    api_key=api_key if embedding_provider == "claude" else None
                )
            except (ImportError, ValueError) as e:
                # Fallback to no embeddings if initialization fails
                print(f"Warning: Could not initialize embeddings: {e}")
                print("Falling back to keyword matching.")
                self.enable_embeddings = False
                self.embedding_model = None

    async def summarize_chunk(
        self,
        chunk_text: str,
        context: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Summarize a chunk using Claude API.

        Args:
            chunk_text: The chunk content (~1000 tokens)
            context: Metadata (file type, sheet/section/page info)

        Returns:
            {
                "summary": "...",  # ~100 tokens
                "tags": ["tag1", "tag2", ...],  # 5-8 keywords
                "token_counts": {"input": ..., "output": ...}
            }
        """
        # Check if Anthropic is available
        if not HAS_ANTHROPIC or Anthropic is None:
            return self._fallback_summary(chunk_text, "Anthropic library not available")

        client = Anthropic(api_key=self.api_key)

        # Build context string
        context_str = self._format_context(context)

        # Construct prompt
        prompt = f"""Summarize this document chunk in approximately {self.summary_tokens} tokens.
Focus on key information, numbers, main points, and important details.
Then extract {self.tag_count} relevant tags (keywords) that capture the content.

Context: {context_str}

Chunk:
{chunk_text}

Format your response as:
SUMMARY: [your concise summary here]
TAGS: [tag1, tag2, tag3, tag4, tag5, tag6, tag7, tag8]
"""

        # Call Claude API
        try:
            response = client.messages.create(
                model=self.model,
                max_tokens=200,  # summary (100) + tags (20) + buffer
                messages=[{
                    "role": "user",
                    "content": prompt
                }]
            )

            # Parse response
            content = response.content[0].text
            summary, tags = self._parse_summary_response(content)

            return {
                "summary": summary,
                "tags": tags,
                "summary_tokens": count_tokens(summary),
                "tag_tokens": count_tokens(" ".join(tags)),
                "api_tokens": {
                    "input": response.usage.input_tokens,
                    "output": response.usage.output_tokens
                }
            }

        except Exception as e:
            # Fallback: simple extraction if API fails
            return self._fallback_summary(chunk_text, str(e))

    def _format_context(self, context: Dict[str, Any]) -> str:
        """Format context metadata into readable string."""
        parts = []
        if "doc_type" in context:
            parts.append(f"Document type: {context['doc_type']}")
        if "sheet_name" in context and context["sheet_name"]:
            parts.append(f"Sheet: {context['sheet_name']}")
        if "section_title" in context and context["section_title"]:
            parts.append(f"Section: {context['section_title']}")
        if "page_range" in context and context["page_range"]:
            parts.append(f"Pages: {context['page_range'][0]}-{context['page_range'][1]}")

        return ", ".join(parts) if parts else "No context"

    def _parse_summary_response(self, content: str) -> Tuple[str, List[str]]:
        """Parse Claude's response to extract summary and tags."""
        summary = ""
        tags = []

        lines = content.strip().split('\n')
        for line in lines:
            if line.startswith("SUMMARY:"):
                summary = line.replace("SUMMARY:", "").strip()
            elif line.startswith("TAGS:"):
                tags_str = line.replace("TAGS:", "").strip()
                # Remove brackets and split
                tags_str = tags_str.replace("[", "").replace("]", "")
                tags = [tag.strip() for tag in tags_str.split(",")]

        # Fallback if parsing failed
        if not summary:
            summary = content[:400]  # First 100 tokens
        if not tags:
            tags = self._extract_keywords(content)

        return summary, tags[:self.tag_count]

    def _extract_keywords(self, text: str) -> List[str]:
        """Simple keyword extraction as fallback."""
        # Remove common words and extract meaningful terms
        words = re.findall(r'\b[a-zA-Z]{4,}\b', text.lower())
        # Get unique words, limit to tag_count
        unique_words = list(dict.fromkeys(words))
        return unique_words[:self.tag_count]

    def _fallback_summary(self, chunk_text: str, error: str) -> Dict[str, Any]:
        """Fallback summary if API call fails."""
        # Take first ~100 tokens
        summary = chunk_text[:400]
        tags = self._extract_keywords(chunk_text)

        return {
            "summary": summary + "...",
            "tags": tags,
            "summary_tokens": count_tokens(summary),
            "tag_tokens": count_tokens(" ".join(tags)),
            "error": f"API call failed: {error}",
            "api_tokens": {"input": 0, "output": 0}
        }

    async def summarize_chunks_parallel(
        self,
        chunks_data: List[Dict[str, Any]],
        batch_size: int = 10
    ) -> List[Dict[str, Any]]:
        """
        Summarize multiple chunks in parallel for better performance.

        Phase 5.6.1: Parallel Chunk Summarization

        Processes chunks in batches to avoid overwhelming the API while achieving
        significant speedup (3-5x) compared to sequential processing.

        Args:
            chunks_data: List of chunk dicts with 'content' and 'context'
            batch_size: Number of chunks to process in parallel (default: 10)

        Returns:
            List of summary results in same order as input
        """
        import asyncio

        results = []

        # Process in batches to avoid overwhelming the API
        for i in range(0, len(chunks_data), batch_size):
            batch = chunks_data[i:i + batch_size]

            # Create tasks for parallel processing
            tasks = []
            for chunk_data in batch:
                task = self.summarize_chunk(
                    chunk_data["content"],
                    chunk_data.get("context", {})
                )
                tasks.append(task)

            # Execute batch in parallel
            batch_results = await asyncio.gather(*tasks, return_exceptions=True)

            # Handle any exceptions
            for j, result in enumerate(batch_results):
                if isinstance(result, Exception):
                    # Use fallback for failed chunks
                    result = self._fallback_summary(
                        batch[j]["content"],
                        str(result)
                    )
                results.append(result)

        return results

    def add_embeddings_to_chunks(self, chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Add semantic embeddings to chunks (Phase 5.2.2).

        Generates embeddings from summary + tags for better semantic matching.

        Args:
            chunks: List of chunks with 'summary' and 'tags' fields

        Returns:
            Same chunks with added 'embedding' field (list of floats)
        """
        if not self.enable_embeddings or not self.embedding_model:
            # No embeddings, return chunks as-is
            return chunks

        if not chunks:
            return chunks

        try:
            # Combine summary and tags for each chunk
            texts_to_embed = []
            for chunk in chunks:
                summary = chunk.get("summary", "")
                tags = chunk.get("tags", [])
                tags_str = ", ".join(tags) if tags else ""
                combined_text = f"{summary} | {tags_str}"
                texts_to_embed.append(combined_text)

            # Batch encode all chunks
            embeddings = self.embedding_model.encode(texts_to_embed)

            # Add embeddings to chunks
            for chunk, embedding in zip(chunks, embeddings):
                # Convert numpy array to list for JSON serialization
                chunk["embedding"] = embedding.tolist()

            return chunks

        except Exception as e:
            # If embedding fails, log warning and return chunks without embeddings
            print(f"Warning: Failed to generate embeddings: {e}")
            return chunks


# ============================================================================
# ChunkCache Class
# ============================================================================

class ChunkCache:
    """Cache document chunks with TTL-based expiration."""

    def __init__(
        self,
        cache_dir: Optional[Path] = None,
        ttl: int = 3600,
        max_size_mb: int = 500
    ):
        """
        Initialize chunk cache.

        Args:
            cache_dir: Cache directory (default: ~/.wyn360/cache/documents/)
            ttl: Time to live in seconds (default: 1 hour)
            max_size_mb: Maximum cache size in MB
        """
        if cache_dir is None:
            cache_dir = Path.home() / ".wyn360" / "cache" / "documents"

        self.cache_dir = cache_dir
        self.ttl = ttl
        self.max_size_mb = max_size_mb

        # Create cache directory
        self.cache_dir.mkdir(parents=True, exist_ok=True)

    def get_file_hash(self, file_path: str) -> str:
        """Generate MD5 hash of file for cache key."""
        hasher = hashlib.md5()

        # Include file path
        hasher.update(file_path.encode())

        # Add file modification time if file exists
        file_obj = Path(file_path)
        if file_obj.exists():
            file_stat = file_obj.stat()
            hasher.update(str(file_stat.st_mtime).encode())

        return hasher.hexdigest()

    def get_cache_path(self, file_hash: str) -> Path:
        """Get cache directory path for a file hash."""
        return self.cache_dir / file_hash

    def load_chunks(self, file_path: str) -> Optional[Dict[str, Any]]:
        """
        Load cached chunks for a file.

        Phase 5.6.2: Supports both compressed (.json.gz) and uncompressed (.json)
        cache files for backward compatibility.

        Args:
            file_path: Path to the document file

        Returns:
            Cached data dict or None if cache miss/expired
        """
        import gzip

        file_hash = self.get_file_hash(file_path)
        cache_path = self.get_cache_path(file_hash)

        # Try compressed files first (new format)
        metadata_file_gz = cache_path / "metadata.json.gz"
        chunks_file_gz = cache_path / "chunks_index.json.gz"

        # Fall back to uncompressed (old format)
        metadata_file = cache_path / "metadata.json"
        chunks_file = cache_path / "chunks_index.json"

        # Determine which files to use
        use_compressed = metadata_file_gz.exists() and chunks_file_gz.exists()
        use_uncompressed = metadata_file.exists() and chunks_file.exists()

        if not use_compressed and not use_uncompressed:
            return None

        # Load metadata
        try:
            if use_compressed:
                # Load compressed metadata
                with gzip.open(metadata_file_gz, 'rt', encoding='utf-8') as f:
                    metadata = json.load(f)
            else:
                # Load uncompressed metadata (backward compatibility)
                with open(metadata_file, 'r') as f:
                    metadata = json.load(f)

            # Check TTL
            age = time.time() - metadata["created_at"]
            if age > metadata["ttl"]:
                # Expired
                self._remove_cache(cache_path)
                return None

            # Load chunks
            if use_compressed:
                # Load compressed chunks
                with gzip.open(chunks_file_gz, 'rt', encoding='utf-8') as f:
                    chunks = json.load(f)
            else:
                # Load uncompressed chunks (backward compatibility)
                with open(chunks_file, 'r') as f:
                    chunks = json.load(f)

            # Phase 5.6.3: Update last_accessed timestamp for LRU eviction
            try:
                metadata["last_accessed"] = time.time()
                # Save updated metadata back with last_accessed time
                if use_compressed:
                    with gzip.open(metadata_file_gz, 'wt', encoding='utf-8') as f:
                        json.dump(metadata, f, indent=2)
                else:
                    with open(metadata_file, 'w') as f:
                        json.dump(metadata, f, indent=2)
            except Exception as access_err:
                # Non-critical - just log and continue
                print(f"Warning: Failed to update last_accessed: {access_err}")

            return {
                "metadata": metadata,
                "chunks": chunks["chunks"]
            }

        except Exception as e:
            print(f"Warning: Failed to load cache: {e}")
            return None

    def save_chunks(
        self,
        file_path: str,
        metadata: DocumentMetadata,
        chunks: List[ChunkMetadata]
    ) -> bool:
        """
        Save chunks to cache.

        Phase 5.6.2: Saves cache files with gzip compression for 50-70%
        storage reduction. Uses .json.gz extension.

        Args:
            file_path: Path to the document file
            metadata: Document metadata
            chunks: List of chunk metadata

        Returns:
            True if saved successfully
        """
        import gzip

        file_hash = self.get_file_hash(file_path)
        cache_path = self.get_cache_path(file_hash)

        # Create cache directory
        cache_path.mkdir(parents=True, exist_ok=True)

        # Check cache size before saving
        self._cleanup_if_needed()

        try:
            # Save metadata with compression
            metadata_file_gz = cache_path / "metadata.json.gz"
            metadata_dict = asdict(metadata)
            # Phase 5.6.3: Initialize last_accessed for LRU eviction
            metadata_dict["last_accessed"] = metadata_dict["created_at"]
            with gzip.open(metadata_file_gz, 'wt', encoding='utf-8') as f:
                json.dump(metadata_dict, f, indent=2)

            # Save chunks with compression
            chunks_file_gz = cache_path / "chunks_index.json.gz"
            chunks_data = {
                "chunks": [asdict(chunk) for chunk in chunks]
            }
            with gzip.open(chunks_file_gz, 'wt', encoding='utf-8') as f:
                json.dump(chunks_data, f, indent=2)

            # Remove old uncompressed files if they exist (migration)
            old_metadata = cache_path / "metadata.json"
            old_chunks = cache_path / "chunks_index.json"
            if old_metadata.exists():
                old_metadata.unlink()
            if old_chunks.exists():
                old_chunks.unlink()

            return True

        except Exception as e:
            print(f"Warning: Failed to save cache: {e}")
            return False

    def clear_cache(self, file_path: Optional[str] = None) -> int:
        """
        Clear cache entries.

        Args:
            file_path: Specific file to clear, or None to clear all

        Returns:
            Number of cache entries removed
        """
        if file_path:
            # Clear specific file
            file_hash = self.get_file_hash(file_path)
            cache_path = self.get_cache_path(file_hash)

            if cache_path.exists():
                self._remove_cache(cache_path)
                return 1
            return 0
        else:
            # Clear all
            count = 0
            for cache_path in self.cache_dir.iterdir():
                if cache_path.is_dir():
                    self._remove_cache(cache_path)
                    count += 1
            return count

    def get_stats(self) -> Dict[str, Any]:
        """Get cache statistics."""
        import gzip  # Phase 5.6.2: Support compressed metadata

        total_files = 0
        total_chunks = 0
        total_size = 0
        oldest_time = None
        newest_time = None
        cache_entries = []

        # Check if cache directory exists
        if not self.cache_dir.exists():
            return {
                "total_files": 0,
                "total_chunks": 0,
                "total_size_mb": 0.0,
                "oldest_cache": "N/A",
                "newest_cache": "N/A",
                "cache_entries": []
            }

        for cache_path in self.cache_dir.iterdir():
            if not cache_path.is_dir():
                continue

            # Phase 5.6.2 & 5.6.3: Support both compressed and uncompressed metadata
            metadata_file_gz = cache_path / "metadata.json.gz"
            metadata_file = cache_path / "metadata.json"

            metadata = None
            if metadata_file_gz.exists():
                # Load compressed metadata
                try:
                    with gzip.open(metadata_file_gz, 'rt', encoding='utf-8') as f:
                        metadata = json.load(f)
                except Exception as e:
                    print(f"Warning: Failed to read compressed metadata: {e}")
            elif metadata_file.exists():
                # Load uncompressed metadata (backward compatibility)
                try:
                    with open(metadata_file, 'r') as f:
                        metadata = json.load(f)
                except Exception as e:
                    print(f"Warning: Failed to read metadata: {e}")

            if metadata is None:
                continue

            try:
                total_files += 1
                total_chunks += metadata["chunk_count"]

                # Calculate directory size
                dir_size = sum(f.stat().st_size for f in cache_path.rglob('*') if f.is_file())
                total_size += dir_size

                # Track times
                created_at = metadata["created_at"]
                if oldest_time is None or created_at < oldest_time:
                    oldest_time = created_at
                if newest_time is None or created_at > newest_time:
                    newest_time = created_at

                # Calculate age
                age_seconds = time.time() - created_at
                age_minutes = int(age_seconds / 60)

                # Phase 5.6.3: Include last_accessed for LRU eviction
                last_accessed = metadata.get("last_accessed", created_at)

                cache_entries.append({
                    "file_path": metadata["file_path"],
                    "chunks": metadata["chunk_count"],
                    "age_seconds": age_seconds,
                    "age_display": self._format_age(age_seconds),
                    "last_accessed": last_accessed  # Phase 5.6.3
                })

            except Exception as e:
                print(f"Warning: Failed to process cache metadata: {e}")

        # Calculate size in MB (use 3 decimal places for better precision with small files)
        size_mb = total_size / (1024 * 1024) if total_size > 0 else 0.0

        return {
            "total_files": total_files,
            "total_chunks": total_chunks,
            "total_size_mb": round(size_mb, 3),
            "oldest_cache": self._format_age(time.time() - oldest_time) if oldest_time else "N/A",
            "newest_cache": self._format_age(time.time() - newest_time) if newest_time else "N/A",
            "cache_entries": sorted(cache_entries, key=lambda x: x["age_seconds"])
        }

    def _format_age(self, seconds: float) -> str:
        """Format age in human-readable format."""
        if seconds < 60:
            return f"{int(seconds)} sec ago"
        elif seconds < 3600:
            return f"{int(seconds / 60)} min ago"
        elif seconds < 86400:
            return f"{int(seconds / 3600)} hours ago"
        else:
            return f"{int(seconds / 86400)} days ago"

    def _remove_cache(self, cache_path: Path):
        """Remove a cache directory."""
        import shutil
        try:
            shutil.rmtree(cache_path)
        except Exception as e:
            print(f"Warning: Failed to remove cache: {e}")

    def _cleanup_if_needed(self):
        """
        Clean up cache if it exceeds max size.

        Phase 5.6.3: Uses LRU (Least Recently Used) eviction strategy instead
        of removing oldest files first. This keeps frequently accessed documents
        cached longer.
        """
        stats = self.get_stats()

        if stats["total_size_mb"] > self.max_size_mb:
            # Phase 5.6.3: Sort by last_accessed (LRU eviction)
            cache_entries = stats["cache_entries"]
            # Sort by last_accessed ascending (least recently used first)
            sorted_entries = sorted(cache_entries, key=lambda x: x["last_accessed"])

            for entry in sorted_entries:  # Least recently used first
                if stats["total_size_mb"] <= self.max_size_mb * 0.9:  # 90% threshold
                    break

                file_path = entry["file_path"]
                file_hash = self.get_file_hash(file_path)
                cache_path = self.get_cache_path(file_hash)

                # Calculate size before removal
                dir_size = sum(f.stat().st_size for f in cache_path.rglob('*') if f.is_file())

                self._remove_cache(cache_path)
                stats["total_size_mb"] -= dir_size / (1024 * 1024)


# ============================================================================
# ChunkRetriever Class
# ============================================================================

class ChunkRetriever:
    """Retrieve relevant chunks based on user queries.

    Phase 5.2.3: Now supports semantic matching with embeddings.
    """

    def __init__(
        self,
        top_k: int = 3,
        embedding_model: Optional[EmbeddingModel] = None,
        similarity_threshold: float = 0.3
    ):
        """
        Initialize retriever.

        Args:
            top_k: Number of top chunks to retrieve (default: 3)
            embedding_model: Optional EmbeddingModel for semantic matching (Phase 5.2.3)
            similarity_threshold: Minimum similarity score for semantic matching (0-1)
        """
        self.top_k = top_k
        self.embedding_model = embedding_model
        self.similarity_threshold = similarity_threshold

    def match_query(
        self,
        query: str,
        chunks: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        Match user query against chunks.

        Phase 5.2.3: Uses semantic matching with embeddings if available,
        falls back to keyword matching otherwise.

        Args:
            query: User's question or search query
            chunks: List of chunk dicts from cache

        Returns:
            Top-K most relevant chunks
        """
        if not chunks:
            return []

        # Check if embeddings are available
        has_embeddings = all("embedding" in chunk for chunk in chunks)

        if has_embeddings and self.embedding_model:
            # SEMANTIC MATCHING (Phase 5.2.3)
            return self._semantic_match(query, chunks)
        else:
            # FALLBACK: KEYWORD MATCHING (Phase 1)
            return self._keyword_match(query, chunks)

    def _semantic_match(
        self,
        query: str,
        chunks: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        Semantic matching using embeddings (Phase 5.2.3).

        Args:
            query: User query
            chunks: List of chunks with embeddings

        Returns:
            Top-K most relevant chunks
        """
        try:
            import numpy as np

            # Encode query
            query_embedding = self.embedding_model.encode(query)

            # Extract chunk embeddings
            chunk_embeddings = np.array([
                chunk["embedding"] for chunk in chunks
            ])

            # Compute similarities
            similarities = self.embedding_model.compute_similarity(
                query_embedding, chunk_embeddings
            )

            # Filter by threshold and rank
            scored_chunks = [
                {
                    "chunk": chunk,
                    "score": float(score),
                    "match_type": "semantic"
                }
                for chunk, score in zip(chunks, similarities)
                if score >= self.similarity_threshold
            ]
            scored_chunks.sort(key=lambda x: x["score"], reverse=True)

            # Return top-K with similarity scores added
            top_chunks = []
            for item in scored_chunks[:self.top_k]:
                chunk = item["chunk"].copy()
                chunk["similarity_score"] = item["score"]
                chunk["match_type"] = "semantic"
                top_chunks.append(chunk)

            return top_chunks

        except Exception as e:
            # If semantic matching fails, fall back to keyword matching
            print(f"Warning: Semantic matching failed: {e}. Falling back to keyword matching.")
            return self._keyword_match(query, chunks)

    def _keyword_match(
        self,
        query: str,
        chunks: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        Keyword matching (fallback when no embeddings).

        Args:
            query: User query
            chunks: List of chunks

        Returns:
            Top-K most relevant chunks
        """
        # Tokenize query (simple word splitting)
        query_terms = self._tokenize_query(query)

        # Score each chunk
        scored_chunks = []
        for chunk in chunks:
            score = self._score_chunk(query_terms, chunk)
            scored_chunks.append({
                "chunk": chunk,
                "score": score
            })

        # Sort by score descending
        scored_chunks.sort(key=lambda x: x["score"], reverse=True)

        # Return top-K with match type
        top_chunks = []
        for item in scored_chunks[:self.top_k]:
            chunk = item["chunk"].copy()
            chunk["match_type"] = "keyword"
            top_chunks.append(chunk)

        return top_chunks

    def _tokenize_query(self, query: str) -> List[str]:
        """Tokenize query into searchable terms."""
        # Convert to lowercase
        query = query.lower()

        # Extract words (4+ characters)
        terms = re.findall(r'\b[a-z0-9]{3,}\b', query)

        return terms

    def _score_chunk(self, query_terms: List[str], chunk: Dict[str, Any]) -> float:
        """
        Score a chunk based on query term overlap with tags.

        Simple keyword matching:
        - Exact tag match: +3 points
        - Partial tag match: +1 point
        - Summary match: +0.5 points

        Args:
            query_terms: List of query terms
            chunk: Chunk dict with tags and summary

        Returns:
            Score (higher = more relevant)
        """
        score = 0.0

        # Get chunk tags and summary
        tags = [tag.lower() for tag in chunk.get("tags", [])]
        summary = chunk.get("summary", "").lower()

        for term in query_terms:
            # Check exact tag match
            if term in tags:
                score += 3.0
                # Skip partial match check for exact matches
                continue

            # Check partial tag match (only if not exact match)
            partial_match = False
            for tag in tags:
                if term in tag or tag in term:
                    score += 1.0
                    partial_match = True
                    break

            # Check summary match (only if no tag match)
            if not partial_match and term in summary:
                score += 0.5

        return score

    def get_relevant_chunks(
        self,
        query: Optional[str],
        chunks: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        Get relevant chunks based on query.

        If no query provided, return all chunks (for full document summary).

        Args:
            query: Optional user query
            chunks: All cached chunks

        Returns:
            Relevant chunks (filtered if query provided)
        """
        if query:
            return self.match_query(query, chunks)
        else:
            # No query: return all chunks (for full document summary)
            return chunks


# ============================================================================
# MultiDocumentRetriever Class (Phase 5.5)
# ============================================================================

class MultiDocumentRetriever:
    """
    Retrieve and search across multiple cached documents.

    Phase 5.5: Multi-Document Queries
    - Unified search across all cached documents
    - Document comparison
    - Cross-document ranking
    """

    def __init__(
        self,
        cache: ChunkCache,
        retriever: Optional[ChunkRetriever] = None,
        embedding_model: Optional[EmbeddingModel] = None
    ):
        """
        Initialize multi-document retriever.

        Args:
            cache: ChunkCache instance for accessing cached documents
            retriever: Optional ChunkRetriever for per-document queries
            embedding_model: Optional EmbeddingModel for semantic matching
        """
        self.cache = cache
        self.retriever = retriever or ChunkRetriever(
            top_k=3,
            embedding_model=embedding_model
        )
        self.embedding_model = embedding_model

    def search_all_documents(
        self,
        query: str,
        top_k: int = 5
    ) -> List[Dict[str, Any]]:
        """
        Search across all cached documents.

        Args:
            query: Search query
            top_k: Number of top results to return across all documents

        Returns:
            List of top-K chunks from all documents, each with:
            - All chunk fields (summary, tags, etc.)
            - file_path: Source document path
            - similarity_score or keyword_score
            - match_type: "semantic" or "keyword"
        """
        all_results = []

        # Get all cached documents
        stats = self.cache.get_stats()

        for entry in stats["cache_entries"]:
            file_path = entry["file_path"]

            # Load chunks for this document
            cached_data = self.cache.load_chunks(file_path)
            if not cached_data:
                continue

            chunks = cached_data["chunks"]

            # Query this document
            matched_chunks = self.retriever.match_query(query, chunks)

            # Add source file path to results
            for chunk in matched_chunks:
                chunk["file_path"] = file_path
                all_results.append(chunk)

        # Sort all results by score
        if all_results and "similarity_score" in all_results[0]:
            # Semantic matching - sort by similarity_score
            all_results.sort(key=lambda x: x.get("similarity_score", 0), reverse=True)
        else:
            # Keyword matching - already sorted by retriever
            pass

        # Return top-K across all documents
        return all_results[:top_k]

    def list_cached_documents(self) -> List[Dict[str, Any]]:
        """
        List all cached documents.

        Returns:
            List of document info dicts with:
            - file_path: Path to document
            - chunks: Number of chunks
            - age_display: Human-readable age
            - doc_type: Document type (excel, pdf, word)
        """
        stats = self.cache.get_stats()

        documents = []
        for entry in stats["cache_entries"]:
            file_path = entry["file_path"]

            # Load metadata
            cached_data = self.cache.load_chunks(file_path)
            if not cached_data:
                continue

            metadata = cached_data["metadata"]

            documents.append({
                "file_path": file_path,
                "chunks": entry["chunks"],
                "age_display": entry["age_display"],
                "doc_type": metadata.get("doc_type", "unknown"),
                "total_tokens": metadata.get("total_tokens", 0)
            })

        return documents

    def compare_documents(
        self,
        file_path1: str,
        file_path2: str,
        aspect: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Compare two cached documents.

        Args:
            file_path1: Path to first document
            file_path2: Path to second document
            aspect: Optional aspect to compare (e.g., "revenue", "expenses")

        Returns:
            Comparison dict with:
            - doc1: Document 1 info and relevant chunks
            - doc2: Document 2 info and relevant chunks
            - comparison: Summary of differences
        """
        # Load both documents
        cached1 = self.cache.load_chunks(file_path1)
        cached2 = self.cache.load_chunks(file_path2)

        if not cached1 or not cached2:
            return {
                "error": "One or both documents not found in cache",
                "doc1_found": cached1 is not None,
                "doc2_found": cached2 is not None
            }

        metadata1 = cached1["metadata"]
        metadata2 = cached2["metadata"]
        chunks1 = cached1["chunks"]
        chunks2 = cached2["chunks"]

        # If aspect specified, filter relevant chunks
        if aspect:
            chunks1 = self.retriever.match_query(aspect, chunks1)
            chunks2 = self.retriever.match_query(aspect, chunks2)

        # Basic comparison metrics
        comparison = {
            "doc1": {
                "file_path": file_path1,
                "doc_type": metadata1.get("doc_type"),
                "total_chunks": metadata1.get("chunk_count"),
                "total_tokens": metadata1.get("total_tokens"),
                "relevant_chunks": chunks1[:3] if aspect else []
            },
            "doc2": {
                "file_path": file_path2,
                "doc_type": metadata2.get("doc_type"),
                "total_chunks": metadata2.get("chunk_count"),
                "total_tokens": metadata2.get("total_tokens"),
                "relevant_chunks": chunks2[:3] if aspect else []
            },
            "comparison": {
                "chunk_count_diff": metadata2.get("chunk_count", 0) - metadata1.get("chunk_count", 0),
                "token_count_diff": metadata2.get("total_tokens", 0) - metadata1.get("total_tokens", 0),
                "same_type": metadata1.get("doc_type") == metadata2.get("doc_type"),
                "aspect_compared": aspect
            }
        }

        return comparison

    def find_cross_references(
        self,
        entity: str,
        min_mentions: int = 1
    ) -> Dict[str, List[Dict[str, Any]]]:
        """
        Find documents that mention a specific entity.

        Args:
            entity: Entity to search for (e.g., "machine learning", "Q1 revenue")
            min_mentions: Minimum number of mentions to include document

        Returns:
            Dict mapping file_path to list of chunks mentioning the entity
        """
        cross_refs = {}

        # Get all cached documents
        stats = self.cache.get_stats()

        for entry in stats["cache_entries"]:
            file_path = entry["file_path"]

            # Load chunks
            cached_data = self.cache.load_chunks(file_path)
            if not cached_data:
                continue

            chunks = cached_data["chunks"]

            # Find chunks mentioning entity
            matching_chunks = self.retriever.match_query(entity, chunks)

            if len(matching_chunks) >= min_mentions:
                cross_refs[file_path] = matching_chunks

        return cross_refs


# ============================================================================
# ExcelReader Class (Phase 2)
# ============================================================================

class ExcelReader:
    """Read and process Excel files with intelligent chunking."""

    def __init__(
        self,
        file_path: str,
        chunk_size: int = 1000,
        include_sheets: Optional[List[str]] = None,
        extract_charts: bool = False,
        extract_named_ranges: bool = True,
        track_formulas: bool = True
    ):
        """
        Initialize Excel reader.

        Args:
            file_path: Path to Excel file
            chunk_size: Target tokens per chunk
            include_sheets: Optional list of sheet names to include
            extract_charts: Extract charts from sheets (Phase 5.4.1)
            extract_named_ranges: Extract named ranges from workbook (Phase 5.4.2)
            track_formulas: Track formula dependencies (Phase 5.4.3)
        """
        self.file_path = Path(file_path)
        self.chunk_size = chunk_size
        self.include_sheets = include_sheets
        self.extract_charts = extract_charts
        self.extract_named_ranges = extract_named_ranges
        self.track_formulas = track_formulas
        self.chunker = DocumentChunker(chunk_size)

    def read(self) -> Dict[str, Any]:
        """
        Read Excel file and return structured data.

        Returns:
            {
                "sheets": [
                    {
                        "name": "Sheet1",
                        "data_region": (min_row, min_col, max_row, max_col),
                        "markdown": "...",
                        "row_count": 100,
                        "col_count": 10,
                        "has_merged_cells": True
                    },
                    ...
                ],
                "total_sheets": 3,
                "total_tokens": 5000
            }
        """
        if not HAS_OPENPYXL:
            raise ImportError(
                "openpyxl not installed. Install with: pip install openpyxl"
            )

        if not self.file_path.exists():
            raise FileNotFoundError(f"Excel file not found: {self.file_path}")

        try:
            # Open workbook (data_only=True to get evaluated formula values)
            workbook = openpyxl.load_workbook(
                self.file_path,
                data_only=True,
                read_only=False
            )

            sheets_data = []
            total_tokens = 0

            # Extract named ranges if enabled (Phase 5.4.2)
            named_ranges = {}
            if self.extract_named_ranges:
                named_ranges = self._extract_named_ranges(workbook)

            # Process each sheet
            for sheet_name in workbook.sheetnames:
                # Filter sheets if include_sheets specified
                if self.include_sheets and sheet_name not in self.include_sheets:
                    continue

                sheet = workbook[sheet_name]

                # Detect data region
                data_region = self._detect_data_region(sheet)

                # Check if sheet has any data
                if data_region is None:
                    continue

                min_row, min_col, max_row, max_col = data_region

                # Get merged cells info
                merged_cells = self._get_merged_cells_map(sheet)

                # Convert to markdown
                markdown = self._sheet_to_markdown(
                    sheet,
                    data_region,
                    merged_cells
                )

                # Count tokens
                tokens = count_tokens(markdown)
                total_tokens += tokens

                sheet_data = {
                    "name": sheet_name,
                    "data_region": data_region,
                    "markdown": markdown,
                    "row_count": max_row - min_row + 1,
                    "col_count": max_col - min_col + 1,
                    "has_merged_cells": len(merged_cells) > 0,
                    "tokens": tokens
                }

                # Extract charts if enabled (Phase 5.4.1)
                if self.extract_charts:
                    charts = self._extract_charts(sheet)
                    if charts:
                        sheet_data["charts"] = charts
                        sheet_data["chart_count"] = len(charts)

                # Track formulas if enabled (Phase 5.4.3)
                if self.track_formulas:
                    formulas = self._track_formulas(sheet)
                    if formulas:
                        sheet_data["formulas"] = formulas
                        sheet_data["formula_count"] = len(formulas)

                sheets_data.append(sheet_data)

            workbook.close()

            result = {
                "sheets": sheets_data,
                "total_sheets": len(sheets_data),
                "total_tokens": total_tokens
            }

            # Add named ranges if any were found (Phase 5.4.2)
            if named_ranges:
                result["named_ranges"] = named_ranges

            return result

        except Exception as e:
            raise Exception(f"Failed to read Excel file: {e}")

    def _detect_data_region(self, sheet) -> Optional[Tuple[int, int, int, int]]:
        """
        Detect the actual data region in a sheet.

        Does not assume data starts at A1. Scans the sheet to find
        the bounding box of all non-empty cells.

        Args:
            sheet: openpyxl worksheet

        Returns:
            (min_row, min_col, max_row, max_col) or None if empty
        """
        min_row = None
        min_col = None
        max_row = None
        max_col = None

        # Scan all cells to find data region
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value is not None:
                    # Update bounding box
                    if min_row is None or cell.row < min_row:
                        min_row = cell.row
                    if max_row is None or cell.row > max_row:
                        max_row = cell.row
                    if min_col is None or cell.column < min_col:
                        min_col = cell.column
                    if max_col is None or cell.column > max_col:
                        max_col = cell.column

        if min_row is None:
            return None  # Empty sheet

        return (min_row, min_col, max_row, max_col)

    def _get_merged_cells_map(self, sheet) -> Dict[Tuple[int, int], Tuple[int, int, int, int]]:
        """
        Get mapping of merged cells.

        Args:
            sheet: openpyxl worksheet

        Returns:
            Dict mapping (row, col) to (min_row, min_col, max_row, max_col)
        """
        merged_map = {}

        for merged_range in sheet.merged_cells.ranges:
            min_row = merged_range.min_row
            min_col = merged_range.min_col
            max_row = merged_range.max_row
            max_col = merged_range.max_col

            # Map all cells in the merged range to the range bounds
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    merged_map[(row, col)] = (min_row, min_col, max_row, max_col)

        return merged_map

    def _sheet_to_markdown(
        self,
        sheet,
        data_region: Tuple[int, int, int, int],
        merged_cells: Dict
    ) -> str:
        """
        Convert sheet data to markdown table.

        Args:
            sheet: openpyxl worksheet
            data_region: (min_row, min_col, max_row, max_col)
            merged_cells: Merged cells mapping

        Returns:
            Markdown formatted table
        """
        min_row, min_col, max_row, max_col = data_region

        lines = []
        lines.append(f"## Sheet: {sheet.title}\n")
        lines.append(f"Data Region: Rows {min_row}-{max_row}, Columns {min_col}-{max_col}\n")

        # Build table
        table_rows = []

        for row_idx in range(min_row, max_row + 1):
            row_cells = []

            for col_idx in range(min_col, max_col + 1):
                cell = sheet.cell(row=row_idx, column=col_idx)
                value = cell.value

                # Handle merged cells
                if (row_idx, col_idx) in merged_cells:
                    merge_info = merged_cells[(row_idx, col_idx)]
                    merge_min_row, merge_min_col, merge_max_row, merge_max_col = merge_info

                    # Only show value in the top-left cell of merged range
                    if row_idx == merge_min_row and col_idx == merge_min_col:
                        # Get the value from the merged cell
                        value = sheet.cell(row=merge_min_row, column=merge_min_col).value
                    else:
                        # Other cells in merged range show empty
                        value = ""

                # Format value
                if value is None:
                    value = ""
                else:
                    value = str(value)

                # Escape pipe characters in cell values
                value = value.replace("|", "\\|")

                row_cells.append(value)

            table_rows.append(row_cells)

        # Format as markdown table
        if table_rows:
            # Header row
            header = "| " + " | ".join(table_rows[0]) + " |"
            lines.append(header)

            # Separator
            num_cols = len(table_rows[0])
            separator = "| " + " | ".join(["---"] * num_cols) + " |"
            lines.append(separator)

            # Data rows
            for row_cells in table_rows[1:]:
                row_line = "| " + " | ".join(row_cells) + " |"
                lines.append(row_line)

        return "\n".join(lines)

    def chunk_sheets(
        self,
        sheets_data: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        Chunk sheet data intelligently.

        Strategy:
        - Each sheet is initially one chunk
        - If sheet exceeds chunk_size, split by row ranges
        - Preserve table structure within chunks

        Args:
            sheets_data: List of sheet data dicts

        Returns:
            List of chunk dicts with metadata
        """
        chunks = []
        chunk_id_counter = 1

        for sheet_data in sheets_data:
            sheet_name = sheet_data["name"]
            markdown = sheet_data["markdown"]
            tokens = sheet_data["tokens"]

            # If sheet fits in one chunk, use as-is
            if tokens <= self.chunk_size:
                chunks.append({
                    "chunk_id": f"{chunk_id_counter:03d}",
                    "sheet_name": sheet_name,
                    "content": markdown,
                    "tokens": tokens,
                    "position": {
                        "sheet": sheet_name,
                        "chunk_type": "full_sheet"
                    }
                })
                chunk_id_counter += 1
            else:
                # Sheet too large - split by row ranges
                # Simple approach: split markdown by lines
                lines = markdown.split("\n")

                # Try to preserve table header + rows in each chunk
                header_lines = []
                data_lines = []
                in_table = False

                for line in lines:
                    if line.startswith("##"):
                        header_lines.append(line)
                    elif line.startswith("|") and "---" in line:
                        # Separator line
                        header_lines.append(line)
                        in_table = True
                    elif line.startswith("|") and not in_table:
                        # Table header
                        header_lines.append(line)
                    elif line.startswith("|") and in_table:
                        # Table data row
                        data_lines.append(line)
                    else:
                        header_lines.append(line)

                # Chunk data rows
                target_chars = self.chunk_size * 4
                current_chunk_lines = header_lines.copy()
                current_chunk_tokens = count_tokens("\n".join(current_chunk_lines))

                for data_line in data_lines:
                    line_tokens = count_tokens(data_line)

                    if current_chunk_tokens + line_tokens > self.chunk_size:
                        # Flush current chunk
                        chunk_content = "\n".join(current_chunk_lines)
                        chunks.append({
                            "chunk_id": f"{chunk_id_counter:03d}",
                            "sheet_name": sheet_name,
                            "content": chunk_content,
                            "tokens": current_chunk_tokens,
                            "position": {
                                "sheet": sheet_name,
                                "chunk_type": "partial"
                            }
                        })
                        chunk_id_counter += 1

                        # Start new chunk with header
                        current_chunk_lines = header_lines.copy()
                        current_chunk_lines.append(data_line)
                        current_chunk_tokens = count_tokens("\n".join(current_chunk_lines))
                    else:
                        current_chunk_lines.append(data_line)
                        current_chunk_tokens += line_tokens

                # Flush final chunk
                if current_chunk_lines:
                    chunk_content = "\n".join(current_chunk_lines)
                    chunks.append({
                        "chunk_id": f"{chunk_id_counter:03d}",
                        "sheet_name": sheet_name,
                        "content": chunk_content,
                        "tokens": current_chunk_tokens,
                        "position": {
                            "sheet": sheet_name,
                            "chunk_type": "partial"
                        }
                    })
                    chunk_id_counter += 1

        return chunks

    def _extract_charts(self, sheet) -> List[Dict[str, Any]]:
        """
        Extract charts from Excel sheet (Phase 5.4.1).

        Args:
            sheet: openpyxl worksheet

        Returns:
            List of chart dictionaries with metadata
        """
        charts = []

        # openpyxl stores charts in sheet._charts
        if hasattr(sheet, '_charts') and sheet._charts:
            for idx, chart in enumerate(sheet._charts):
                chart_info = {
                    "chart_id": f"chart_{idx + 1}",
                    "type": type(chart).__name__,  # BarChart, LineChart, PieChart, etc.
                    "title": getattr(chart, 'title', None),
                    "anchor": str(chart.anchor) if hasattr(chart, 'anchor') else None,
                }

                # Extract series information if available
                if hasattr(chart, 'series') and chart.series:
                    chart_info["series_count"] = len(chart.series)

                charts.append(chart_info)

        return charts

    def _extract_named_ranges(self, workbook) -> Dict[str, Dict[str, Any]]:
        """
        Extract named ranges from workbook (Phase 5.4.2).

        Args:
            workbook: openpyxl workbook

        Returns:
            Dictionary of named ranges {name: {scope, refers_to, ...}}
        """
        named_ranges = {}

        # openpyxl stores defined names in workbook.defined_names
        if hasattr(workbook, 'defined_names'):
            for name, definition in workbook.defined_names.items():
                try:
                    named_ranges[name] = {
                        "name": name,
                        "refers_to": str(definition.attr_text) if hasattr(definition, 'attr_text') else str(definition),
                        "scope": "workbook"  # Could be sheet-specific
                    }
                except Exception:
                    # Skip if definition can't be parsed
                    continue

        return named_ranges

    def _track_formulas(self, sheet) -> List[Dict[str, Any]]:
        """
        Track formula cells in sheet (Phase 5.4.3).

        Args:
            sheet: openpyxl worksheet

        Returns:
            List of formula dictionaries {cell, formula, value, dependencies}
        """
        formulas = []

        for row in sheet.iter_rows():
            for cell in row:
                # Check if cell contains a formula
                if cell.value and isinstance(cell.value, str) and cell.value.startswith('='):
                    formula_info = {
                        "cell": cell.coordinate,
                        "formula": cell.value,
                        "sheet": sheet.title
                    }

                    # Try to get the evaluated value
                    try:
                        # Open workbook with data_only=True to get evaluated values
                        formula_info["evaluated_value"] = cell.value
                    except Exception:
                        formula_info["evaluated_value"] = None

                    formulas.append(formula_info)

        return formulas


# ============================================================================
# WordReader Class (Phase 3)
# ============================================================================

class WordReader:
    """Read and process Word documents with intelligent chunking."""

    def __init__(
        self,
        file_path: str,
        chunk_size: int = 1000,
        image_handling: str = "describe"
    ):
        """
        Initialize Word reader.

        Args:
            file_path: Path to Word file
            chunk_size: Target tokens per chunk
            image_handling: How to handle images (skip|describe|vision)
        """
        self.file_path = Path(file_path)
        self.chunk_size = chunk_size
        self.image_handling = image_handling
        self.chunker = DocumentChunker(chunk_size)

    async def read(self, image_processor: Optional['ImageProcessor'] = None) -> Dict[str, Any]:
        """
        Read Word file and return structured data.

        Args:
            image_processor: Optional ImageProcessor for vision mode

        Returns:
            {
                "sections": [
                    {
                        "level": 1,
                        "title": "Introduction",
                        "content": "...",  # markdown
                        "tokens": 500
                    },
                    ...
                ],
                "total_sections": 5,
                "total_tokens": 5000,
                "has_tables": True,
                "has_images": True,
                "images": [...],  # NEW: list of image descriptions if vision mode
                "vision_tokens_used": 0  # NEW: tokens used for vision API
            }
        """
        if not HAS_PYTHON_DOCX:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            )

        if not self.file_path.exists():
            raise FileNotFoundError(f"Word file not found: {self.file_path}")

        try:
            # Open document
            doc = Document(str(self.file_path))

            # Extract images if vision mode enabled
            image_descriptions = []
            vision_tokens_used = 0

            if image_processor and self.image_handling == "vision":
                images = self._extract_images(doc)
                if images:
                    # Process images with vision API
                    image_descriptions = await self._process_images_with_vision(
                        images=images,
                        image_processor=image_processor,
                        doc_context={"doc_type": "word", "file_name": self.file_path.name}
                    )
                    # Track vision tokens
                    vision_tokens_used = sum(img.get("tokens_used", 0) for img in image_descriptions)

            # Extract sections with structure
            sections = self._extract_sections(doc, image_descriptions)

            # Calculate totals
            total_tokens = sum(s["tokens"] for s in sections)
            has_tables = any(s.get("has_tables", False) for s in sections)
            has_images = any(s.get("has_images", False) for s in sections)

            result = {
                "sections": sections,
                "total_sections": len(sections),
                "total_tokens": total_tokens,
                "has_tables": has_tables,
                "has_images": has_images,
                "vision_tokens_used": vision_tokens_used
            }

            # Add image descriptions if vision mode was used
            if image_descriptions:
                result["images"] = image_descriptions

            return result

        except Exception as e:
            raise Exception(f"Failed to read Word file: {e}")

    def _extract_sections(self, doc, image_descriptions: List[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
        """
        Extract document structure with headings as section boundaries.

        Args:
            doc: python-docx Document object
            image_descriptions: Optional list of image descriptions from vision API

        Returns:
            List of section dicts with title, content, level
        """
        if image_descriptions is None:
            image_descriptions = []
        sections = []
        current_section = None
        current_content = []

        for element in doc.element.body:
            # Check if paragraph
            if element.tag.endswith('p'):
                para = None
                # Find the paragraph object
                for p in doc.paragraphs:
                    if p._element == element:
                        para = p
                        break

                if para is None:
                    continue

                # Check if heading
                if para.style.name.startswith('Heading'):
                    # Flush current section
                    if current_section is not None:
                        content_md = "\n\n".join(current_content)
                        current_section["content"] = content_md
                        current_section["tokens"] = count_tokens(content_md)
                        sections.append(current_section)

                    # Start new section
                    level = int(para.style.name.replace('Heading ', '').replace('Heading', '1'))
                    current_section = {
                        "level": level,
                        "title": para.text,
                        "has_tables": False,
                        "has_images": False
                    }
                    current_content = []
                else:
                    # Regular paragraph
                    if para.text.strip():
                        current_content.append(para.text)

            # Check if table
            elif element.tag.endswith('tbl'):
                # Find the table object
                table = None
                for t in doc.tables:
                    if t._element == element:
                        table = t
                        break

                if table is not None:
                    # Convert table to markdown
                    table_md = self._table_to_markdown(table)
                    current_content.append(table_md)

                    if current_section:
                        current_section["has_tables"] = True

        # Flush final section
        if current_section is not None:
            content_md = "\n\n".join(current_content)
            current_section["content"] = content_md
            current_section["tokens"] = count_tokens(content_md)
            sections.append(current_section)

        # If no sections found, create a single section with all content
        if not sections and current_content:
            content_md = "\n\n".join(current_content)
            sections.append({
                "level": 1,
                "title": "Document",
                "content": content_md,
                "tokens": count_tokens(content_md),
                "has_tables": False,
                "has_images": False
            })

        return sections

    def _table_to_markdown(self, table) -> str:
        """
        Convert Word table to markdown.

        Args:
            table: python-docx Table object

        Returns:
            Markdown formatted table
        """
        lines = []

        # Get all rows
        rows = list(table.rows)
        if not rows:
            return ""

        # Header row (first row)
        header_cells = [cell.text.strip() for cell in rows[0].cells]
        header_line = "| " + " | ".join(header_cells) + " |"
        lines.append(header_line)

        # Separator
        separator = "| " + " | ".join(["---"] * len(header_cells)) + " |"
        lines.append(separator)

        # Data rows
        for row in rows[1:]:
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            row_line = "| " + " | ".join(cells) + " |"
            lines.append(row_line)

        return "\n".join(lines)

    def _extract_images(self, doc) -> List[Dict[str, Any]]:
        """
        Extract all images from Word document.

        Args:
            doc: python-docx Document object

        Returns:
            List of image dicts with {
                "data": bytes,
                "format": "png"|"jpeg"|"gif",
                "context": {"index": N, "doc_type": "word"}
            }
        """
        images = []

        try:
            # Method 1: Inline shapes (document.inline_shapes)
            for idx, shape in enumerate(doc.inline_shapes):
                try:
                    # Get image bytes through the relationship
                    image_rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
                    image_part = doc.part.related_parts[image_rId]
                    image_data = image_part.blob

                    # Detect format from content type
                    content_type = image_part.content_type  # 'image/png', 'image/jpeg', etc.
                    image_format = content_type.split('/')[-1] if '/' in content_type else 'png'

                    images.append({
                        "data": image_data,
                        "format": image_format,
                        "context": {
                            "doc_type": "word",
                            "index": idx,
                            "shape_type": "inline"
                        }
                    })
                except Exception:
                    # Skip images that can't be extracted
                    continue

        except Exception:
            # If inline_shapes fails, continue
            pass

        return images

    async def _process_images_with_vision(
        self,
        images: List[Dict[str, Any]],
        image_processor: 'ImageProcessor',
        doc_context: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """
        Process images with vision API.

        Args:
            images: List of image data dicts
            image_processor: ImageProcessor instance
            doc_context: Context about the document

        Returns:
            List of image description dicts with {
                "description": "...",
                "image_type": "chart"|"diagram"|...,
                "tokens_used": N,
                "index": N
            }
        """
        if not images:
            return []

        # Add document context to each image
        for img in images:
            img["context"].update(doc_context)

        # Batch process images
        descriptions = await image_processor.describe_images_batch(
            images=images,
            context=doc_context
        )

        # Add index to each description for reference
        for idx, desc in enumerate(descriptions):
            desc["index"] = images[idx]["context"]["index"]

        return descriptions

    def chunk_sections(
        self,
        sections: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        Chunk sections intelligently.

        Strategy:
        - Try to keep complete sections in chunks
        - Split large sections by paragraphs
        - Preserve section hierarchy in metadata

        Args:
            sections: List of section dicts

        Returns:
            List of chunk dicts with metadata
        """
        chunks = []
        chunk_id_counter = 1

        for section in sections:
            title = section["title"]
            content = section["content"]
            tokens = section["tokens"]
            level = section["level"]

            # If section fits in one chunk, use as-is
            if tokens <= self.chunk_size:
                chunks.append({
                    "chunk_id": f"{chunk_id_counter:03d}",
                    "section_title": title,
                    "section_level": level,
                    "content": content,
                    "tokens": tokens,
                    "position": {
                        "section": title,
                        "level": level,
                        "chunk_type": "full_section"
                    }
                })
                chunk_id_counter += 1
            else:
                # Section too large - split by paragraphs
                paragraphs = content.split("\n\n")

                current_chunk_paras = []
                current_chunk_tokens = 0

                for para in paragraphs:
                    para_tokens = count_tokens(para)

                    if current_chunk_tokens + para_tokens > self.chunk_size and current_chunk_paras:
                        # Flush current chunk
                        chunk_content = "\n\n".join(current_chunk_paras)
                        chunks.append({
                            "chunk_id": f"{chunk_id_counter:03d}",
                            "section_title": title,
                            "section_level": level,
                            "content": chunk_content,
                            "tokens": current_chunk_tokens,
                            "position": {
                                "section": title,
                                "level": level,
                                "chunk_type": "partial"
                            }
                        })
                        chunk_id_counter += 1

                        # Start new chunk
                        current_chunk_paras = [para]
                        current_chunk_tokens = para_tokens
                    else:
                        current_chunk_paras.append(para)
                        current_chunk_tokens += para_tokens

                # Flush final chunk
                if current_chunk_paras:
                    chunk_content = "\n\n".join(current_chunk_paras)
                    chunks.append({
                        "chunk_id": f"{chunk_id_counter:03d}",
                        "section_title": title,
                        "section_level": level,
                        "content": chunk_content,
                        "tokens": current_chunk_tokens,
                        "position": {
                            "section": title,
                            "level": level,
                            "chunk_type": "partial"
                        }
                    })
                    chunk_id_counter += 1

        return chunks


# ============================================================================
# PDF Reader (Phase 4)
# ============================================================================

class PDFReader:
    """
    Read PDF files with page-aware chunking.

    Features:
    - Support for both pymupdf (default, fast) and pdfplumber (better tables)
    - Page-by-page text extraction
    - Table detection and extraction
    - Multi-column layout handling
    - Section detection via font size analysis
    - Page-aware chunking (3-5 pages per chunk)
    - Page range filtering

    Usage:
        reader = PDFReader(file_path="document.pdf", engine="pymupdf")
        result = reader.read(page_range=(10, 20))
        chunks = reader.chunk_pages(result["pages"])
    """

    def __init__(
        self,
        file_path: str,
        chunk_size: int = 1000,
        engine: str = "pymupdf",
        pages_per_chunk: int = 3,
        image_handling: str = "describe",
        enable_ocr: bool = False,
        ocr_language: str = "eng"
    ):
        """
        Initialize PDF reader.

        Args:
            file_path: Path to PDF file
            chunk_size: Target tokens per chunk (used as guideline)
            engine: PDF engine to use ("pymupdf" or "pdfplumber")
            pages_per_chunk: Target pages per chunk (3-5 recommended)
            image_handling: How to handle images ("skip", "describe", "vision")
            enable_ocr: Enable OCR for scanned PDFs (Phase 5.3)
            ocr_language: Tesseract language code (eng, spa, fra, etc.)
        """
        self.file_path = Path(file_path)
        self.chunk_size = chunk_size
        self.engine = engine.lower()
        self.pages_per_chunk = pages_per_chunk
        self.image_handling = image_handling
        self.enable_ocr = enable_ocr
        self.ocr_language = ocr_language
        self.chunker = DocumentChunker(chunk_size)

        # Validate engine
        if self.engine not in ["pymupdf", "pdfplumber"]:
            raise ValueError(f"Unknown PDF engine: {engine}. Use 'pymupdf' or 'pdfplumber'")

    async def read(
        self,
        page_range: Optional[Tuple[int, int]] = None,
        image_processor: Optional['ImageProcessor'] = None
    ) -> Dict[str, Any]:
        """
        Read PDF file and extract pages.

        Args:
            page_range: Optional (start_page, end_page) tuple (1-indexed)
            image_processor: Optional ImageProcessor for vision mode

        Returns:
            Dictionary with:
                - pages: List of page dictionaries
                - total_pages: Total page count
                - page_range_read: Actual range read
                - total_tokens: Total token count
                - has_tables: Whether tables were detected
                - images: List of image descriptions if vision mode enabled
                - vision_tokens_used: Tokens used for vision API calls

        Raises:
            ImportError: If required PDF library not installed
            FileNotFoundError: If file doesn't exist
        """
        # Check dependencies
        if self.engine == "pymupdf" and not HAS_PYMUPDF:
            raise ImportError(
                "pymupdf not installed. Install with: pip install pymupdf"
            )
        elif self.engine == "pdfplumber" and not HAS_PDFPLUMBER:
            raise ImportError(
                "pdfplumber not installed. Install with: pip install pdfplumber"
            )

        # Check file exists
        if not self.file_path.exists():
            raise FileNotFoundError(f"PDF file not found: {self.file_path}")

        # Read using appropriate engine
        if self.engine == "pymupdf":
            return await self._read_with_pymupdf(page_range, image_processor)
        else:
            return await self._read_with_pdfplumber(page_range, image_processor)

    async def _read_with_pymupdf(
        self,
        page_range: Optional[Tuple[int, int]] = None,
        image_processor: Optional['ImageProcessor'] = None
    ) -> Dict[str, Any]:
        """Read PDF using PyMuPDF (faster, general-purpose)."""
        doc = pymupdf.open(str(self.file_path))
        total_pages = len(doc)

        # Determine page range
        if page_range:
            start_page = max(1, page_range[0])
            end_page = min(total_pages, page_range[1])
        else:
            start_page = 1
            end_page = total_pages

        pages = []
        total_tokens = 0
        has_tables = False
        all_images = []

        # Initialize OCR processor if enabled (Phase 5.3.2)
        ocr_processor = None
        if self.enable_ocr and HAS_PYTESSERACT:
            try:
                ocr_processor = OCRProcessor(language=self.ocr_language)
            except RuntimeError:
                # Tesseract not installed, skip OCR
                pass

        # Extract pages
        for page_num in range(start_page - 1, end_page):
            page = doc[page_num]

            # Extract text
            text = page.get_text()

            # Check if page is scanned and use OCR if needed (Phase 5.3.2)
            ocr_used = False
            ocr_confidence = None
            if ocr_processor and len(text.strip()) < 50:
                # Page might be scanned, render as image and check
                try:
                    from PIL import Image
                    pix = page.get_pixmap(dpi=300)
                    image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                    # Check if truly scanned
                    if ocr_processor.is_scanned_page(image, text):
                        # Extract text with OCR
                        ocr_result = ocr_processor.extract_text(image)
                        text = ocr_result["text"]
                        ocr_used = True
                        ocr_confidence = ocr_result["confidence"]
                except Exception:
                    # OCR failed, continue with original text
                    pass

            # Try to detect tables (basic detection)
            tables = []
            try:
                # PyMuPDF table detection (if available)
                page_tables = page.find_tables()
                if page_tables:
                    has_tables = True
                    for table in page_tables.tables:
                        table_data = table.extract()
                        if table_data:
                            markdown = self._table_to_markdown_pymupdf(table_data)
                            tables.append(markdown)
            except (AttributeError, Exception):
                # Table extraction not available or failed
                pass

            # Extract images if vision mode enabled
            if image_processor and self.image_handling == "vision":
                page_images = self._extract_images_pymupdf(page, page_num + 1)
                all_images.extend(page_images)

            # Combine text and tables
            content = text
            if tables:
                content += "\n\n" + "\n\n".join(tables)

            tokens = count_tokens(content)
            total_tokens += tokens

            page_data = {
                "page_number": page_num + 1,
                "content": content,
                "tokens": tokens,
                "has_tables": bool(tables)
            }

            # Add OCR metadata if used (Phase 5.3.2)
            if ocr_used:
                page_data["ocr_used"] = True
                page_data["ocr_confidence"] = ocr_confidence

            pages.append(page_data)

        doc.close()

        # Process images with vision API if enabled
        image_descriptions = []
        vision_tokens_used = 0

        if all_images and image_processor and self.image_handling == "vision":
            image_descriptions = await self._process_images_with_vision(
                images=all_images,
                image_processor=image_processor,
                doc_context={"doc_type": "pdf", "file_name": self.file_path.name}
            )
            vision_tokens_used = sum(img.get("tokens_used", 0) for img in image_descriptions)

        result = {
            "pages": pages,
            "total_pages": total_pages,
            "page_range_read": (start_page, end_page),
            "total_tokens": total_tokens,
            "has_tables": has_tables,
            "vision_tokens_used": vision_tokens_used
        }

        if image_descriptions:
            result["images"] = image_descriptions

        return result

    async def _read_with_pdfplumber(
        self,
        page_range: Optional[Tuple[int, int]] = None,
        image_processor: Optional['ImageProcessor'] = None
    ) -> Dict[str, Any]:
        """Read PDF using pdfplumber (better for complex tables)."""
        all_images = []

        # Initialize OCR processor if enabled (Phase 5.3.2)
        ocr_processor = None
        if self.enable_ocr and HAS_PYTESSERACT and HAS_PDF2IMAGE:
            try:
                ocr_processor = OCRProcessor(language=self.ocr_language)
            except RuntimeError:
                # Tesseract not installed, skip OCR
                pass

        with pdfplumber.open(str(self.file_path)) as pdf:
            total_pages = len(pdf.pages)

            # Determine page range
            if page_range:
                start_page = max(1, page_range[0])
                end_page = min(total_pages, page_range[1])
            else:
                start_page = 1
                end_page = total_pages

            pages = []
            total_tokens = 0
            has_tables = False

            # Extract pages
            for page_num in range(start_page - 1, end_page):
                page = pdf.pages[page_num]

                # Extract text
                text = page.extract_text() or ""

                # Check if page is scanned and use OCR if needed (Phase 5.3.2)
                ocr_used = False
                ocr_confidence = None
                if ocr_processor and len(text.strip()) < 50:
                    # Page might be scanned, render as image
                    try:
                        from pdf2image import convert_from_path
                        from PIL import Image

                        # Convert single page to image
                        images = convert_from_path(
                            str(self.file_path),
                            first_page=page_num + 1,
                            last_page=page_num + 1,
                            dpi=300
                        )

                        if images:
                            image = images[0]
                            # Check if truly scanned
                            if ocr_processor.is_scanned_page(image, text):
                                # Extract text with OCR
                                ocr_result = ocr_processor.extract_text(image)
                                text = ocr_result["text"]
                                ocr_used = True
                                ocr_confidence = ocr_result["confidence"]
                    except Exception:
                        # OCR failed, continue with original text
                        pass

                # Extract tables
                tables = []
                page_tables = page.extract_tables()
                if page_tables:
                    has_tables = True
                    for table_data in page_tables:
                        if table_data:
                            markdown = self._table_to_markdown_pdfplumber(table_data)
                            tables.append(markdown)

                # Extract images if vision mode enabled
                if image_processor and self.image_handling == "vision":
                    page_images = self._extract_images_pdfplumber(page, page_num + 1)
                    all_images.extend(page_images)

                # Combine text and tables
                content = text
                if tables:
                    content += "\n\n" + "\n\n".join(tables)

                tokens = count_tokens(content)
                total_tokens += tokens

                page_data = {
                    "page_number": page_num + 1,
                    "content": content,
                    "tokens": tokens,
                    "has_tables": bool(tables)
                }

                # Add OCR metadata if used (Phase 5.3.2)
                if ocr_used:
                    page_data["ocr_used"] = True
                    page_data["ocr_confidence"] = ocr_confidence

                pages.append(page_data)

        # Process images with vision API if enabled
        image_descriptions = []
        vision_tokens_used = 0

        if all_images and image_processor and self.image_handling == "vision":
            image_descriptions = await self._process_images_with_vision(
                images=all_images,
                image_processor=image_processor,
                doc_context={"doc_type": "pdf", "file_name": self.file_path.name}
            )
            vision_tokens_used = sum(img.get("tokens_used", 0) for img in image_descriptions)

        result = {
            "pages": pages,
            "total_pages": total_pages,
            "page_range_read": (start_page, end_page),
            "total_tokens": total_tokens,
            "has_tables": has_tables,
            "vision_tokens_used": vision_tokens_used
        }

        if image_descriptions:
            result["images"] = image_descriptions

        return result

    def _table_to_markdown_pymupdf(self, table_data: List[List]) -> str:
        """Convert PyMuPDF table data to markdown."""
        if not table_data:
            return ""

        lines = []

        # Header row
        if table_data:
            header = table_data[0]
            header_cells = [str(cell or "").replace("|", "\\|").strip() for cell in header]
            lines.append("| " + " | ".join(header_cells) + " |")
            lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

        # Data rows
        for row in table_data[1:]:
            cells = [str(cell or "").replace("|", "\\|").strip() for cell in row]
            lines.append("| " + " | ".join(cells) + " |")

        return "\n".join(lines)

    def _table_to_markdown_pdfplumber(self, table_data: List[List]) -> str:
        """Convert pdfplumber table data to markdown."""
        if not table_data:
            return ""

        lines = []

        # Header row
        if table_data:
            header = table_data[0]
            header_cells = [str(cell or "").replace("|", "\\|").strip() for cell in header]
            lines.append("| " + " | ".join(header_cells) + " |")
            lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

        # Data rows
        for row in table_data[1:]:
            cells = [str(cell or "").replace("|", "\\|").strip() for cell in row]
            lines.append("| " + " | ".join(cells) + " |")

        return "\n".join(lines)

    def _extract_images_pymupdf(self, page, page_num: int) -> List[Dict[str, Any]]:
        """
        Extract images from PDF page using PyMuPDF.

        Args:
            page: PyMuPDF page object
            page_num: Page number (1-indexed)

        Returns:
            List of image dicts with data, format, and context
        """
        images = []

        try:
            image_list = page.get_images(full=True)

            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]  # image xref
                    base_image = page.parent.extract_image(xref)

                    image_data = base_image["image"]
                    image_ext = base_image["ext"]  # png, jpeg, etc.

                    images.append({
                        "data": image_data,
                        "format": image_ext,
                        "context": {
                            "doc_type": "pdf",
                            "page_number": page_num,
                            "index": img_index
                        }
                    })
                except Exception:
                    # Skip if individual image extraction fails
                    continue
        except Exception:
            # Skip if page image extraction fails
            pass

        return images

    def _extract_images_pdfplumber(self, page, page_num: int) -> List[Dict[str, Any]]:
        """
        Extract images from PDF page using pdfplumber.

        Note: pdfplumber has limited image extraction capabilities.
        Images are detected but raw data extraction requires additional processing.

        Args:
            page: pdfplumber page object
            page_num: Page number (1-indexed)

        Returns:
            List of image dicts (note: may have limited data)
        """
        images = []

        try:
            # pdfplumber provides image objects but not raw data directly
            page_images = page.images

            for idx, img_info in enumerate(page_images):
                # pdfplumber doesn't provide easy access to raw image bytes
                # We'll create placeholder entries for now
                images.append({
                    "data": None,  # pdfplumber limitation
                    "format": "unknown",
                    "context": {
                        "doc_type": "pdf",
                        "page_number": page_num,
                        "index": idx,
                        "note": "pdfplumber image extraction limited",
                        "bbox": img_info.get("bbox") if hasattr(img_info, "get") else None
                    }
                })
        except Exception:
            pass

        return images

    async def _process_images_with_vision(
        self,
        images: List[Dict[str, Any]],
        image_processor: 'ImageProcessor',
        doc_context: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """
        Process images with vision API.

        Args:
            images: List of image dicts from extraction
            image_processor: ImageProcessor instance
            doc_context: Document context (file_name, doc_type)

        Returns:
            List of image description dicts
        """
        if not images:
            return []

        # Filter out images without data (e.g., from pdfplumber)
        valid_images = [img for img in images if img.get("data") is not None]

        if not valid_images:
            return []

        # Add document context to each image
        for img in valid_images:
            img["context"].update(doc_context)

        # Batch process images
        descriptions = await image_processor.describe_images_batch(
            images=valid_images,
            context=doc_context
        )

        # Add index and page number to each description for reference
        for idx, desc in enumerate(descriptions):
            if idx < len(valid_images):
                desc["page_number"] = valid_images[idx]["context"]["page_number"]
                desc["index"] = valid_images[idx]["context"]["index"]

        return descriptions

    def chunk_pages(self, pages: List[Dict]) -> List[Dict]:
        """
        Chunk pages with page-aware strategy.

        Strategy:
        1. Group pages into chunks (default: 3-5 pages per chunk)
        2. If a chunk exceeds chunk_size tokens, split it further
        3. Preserve page boundaries in metadata

        Args:
            pages: List of page dictionaries from read()

        Returns:
            List of chunk dictionaries with:
                - chunk_id: Unique chunk identifier
                - page_range: (start_page, end_page) for this chunk
                - content: Combined content
                - tokens: Token count
                - has_tables: Whether chunk contains tables
                - position: Chunk position metadata
        """
        if not pages:
            return []

        chunks = []
        chunk_id_counter = 1

        i = 0
        while i < len(pages):
            # Collect pages for this chunk
            chunk_pages = []
            chunk_tokens = 0
            chunk_has_tables = False

            # Try to fit pages_per_chunk pages, but stop if we exceed chunk_size
            for j in range(self.pages_per_chunk):
                if i + j >= len(pages):
                    break

                page = pages[i + j]
                page_tokens = page["tokens"]

                # If adding this page would exceed chunk_size and we already have at least 1 page,
                # stop here (unless it's the first page)
                if chunk_tokens > 0 and chunk_tokens + page_tokens > self.chunk_size:
                    break

                chunk_pages.append(page)
                chunk_tokens += page_tokens
                if page.get("has_tables"):
                    chunk_has_tables = True

            # If no pages were added (single page too large), take it anyway
            if not chunk_pages and i < len(pages):
                chunk_pages.append(pages[i])
                chunk_tokens = pages[i]["tokens"]
                chunk_has_tables = pages[i].get("has_tables", False)

            # Create chunk
            if chunk_pages:
                start_page = chunk_pages[0]["page_number"]
                end_page = chunk_pages[-1]["page_number"]

                # Combine content from all pages
                content_parts = []
                for page in chunk_pages:
                    content_parts.append(
                        f"[Page {page['page_number']}]\n{page['content']}"
                    )

                chunks.append({
                    "chunk_id": f"{chunk_id_counter:03d}",
                    "page_range": (start_page, end_page),
                    "content": "\n\n".join(content_parts),
                    "tokens": chunk_tokens,
                    "has_tables": chunk_has_tables,
                    "position": {
                        "start_page": start_page,
                        "end_page": end_page,
                        "chunk_type": "page_range"
                    }
                })
                chunk_id_counter += 1
                i += len(chunk_pages)
            else:
                # Shouldn't happen, but safety check
                i += 1

        return chunks
