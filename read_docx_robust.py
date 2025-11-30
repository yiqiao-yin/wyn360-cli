#!/usr/bin/env python3
"""
Robust script to read a Word document with error handling for complex formatting.
"""

from docx import Document
import sys
import os

def read_word_document_robust(file_path):
    """Read a Word document with robust error handling."""
    try:
        # Check if file exists
        if not os.path.exists(file_path):
            print(f"Error: File not found: {file_path}")
            return None
        
        print(f"File size: {os.path.getsize(file_path)} bytes")
        
        # Load the document
        doc = Document(file_path)
        
        # Extract basic document properties
        core_props = doc.core_properties
        print(f"Title: {core_props.title or 'N/A'}")
        print(f"Author: {core_props.author or 'N/A'}")
        print(f"Subject: {core_props.subject or 'N/A'}")
        print(f"Created: {core_props.created or 'N/A'}")
        print(f"Modified: {core_props.modified or 'N/A'}")
        
        # Extract text from paragraphs with error handling
        full_text = []
        paragraph_count = 0
        
        print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
        
        for i, paragraph in enumerate(doc.paragraphs):
            try:
                if paragraph.text and paragraph.text.strip():
                    full_text.append(f"[Para {i+1}] {paragraph.text.strip()}")
                    paragraph_count += 1
            except Exception as e:
                print(f"Error reading paragraph {i+1}: {e}")
                continue
        
        print(f"Successfully read paragraphs: {paragraph_count}")
        
        # Try to extract table content with error handling
        table_count = 0
        print(f"Total tables: {len(doc.tables)}")
        
        for i, table in enumerate(doc.tables):
            try:
                full_text.append(f"\n[Table {i+1}]")
                for row_idx, row in enumerate(table.rows):
                    try:
                        row_text = []
                        for cell_idx, cell in enumerate(row.cells):
                            try:
                                if cell.text and cell.text.strip():
                                    row_text.append(cell.text.strip())
                            except Exception as e:
                                print(f"Error reading table {i+1}, row {row_idx+1}, cell {cell_idx+1}: {e}")
                                row_text.append("[Error reading cell]")
                        if row_text:
                            full_text.append(" | ".join(row_text))
                    except Exception as e:
                        print(f"Error reading table {i+1}, row {row_idx+1}: {e}")
                        continue
                table_count += 1
            except Exception as e:
                print(f"Error reading table {i+1}: {e}")
                continue
        
        print(f"Successfully read tables: {table_count}")
        
        return "\n".join(full_text) if full_text else "No readable content found."
        
    except Exception as e:
        print(f"Error reading document: {str(e)}")
        return None

if __name__ == "__main__":
    file_path = "/home/workbench/wyn360-cli/wyn360-cli/tests/test_files/test_cl.docx"
    
    print("Reading Word document with robust error handling...")
    content = read_word_document_robust(file_path)
    
    if content:
        print("\n" + "=" * 60)
        print("DOCUMENT CONTENT:")
        print("=" * 60)
        print(content)
        print("=" * 60)
        print(f"\nTotal content length: {len(content)} characters")
    else:
        print("Failed to read document.")