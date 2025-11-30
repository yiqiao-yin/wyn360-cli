#!/usr/bin/env python3
"""
Simple script to read a Word document and extract its text content.
"""

try:
    from docx import Document
    import sys
    import os
    
    def read_word_document(file_path):
        """Read a Word document and return its text content."""
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                print(f"Error: File not found: {file_path}")
                return None
            
            # Load the document
            doc = Document(file_path)
            
            # Extract text from all paragraphs
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only include non-empty paragraphs
                    full_text.append(paragraph.text.strip())
            
            # Extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            return "\n\n".join(full_text)
            
        except Exception as e:
            print(f"Error reading document: {str(e)}")
            return None
    
    if __name__ == "__main__":
        file_path = "/home/workbench/wyn360-cli/wyn360-cli/tests/test_files/test_cl.docx"
        
        print("Reading Word document...")
        content = read_word_document(file_path)
        
        if content:
            print("Document Content:")
            print("=" * 50)
            print(content)
            print("=" * 50)
            print(f"\nDocument successfully read. Total characters: {len(content)}")
        else:
            print("Failed to read document.")

except ImportError:
    print("Error: python-docx is not installed.")
    print("Please install it with: pip install python-docx")
except Exception as e:
    print(f"Unexpected error: {str(e)}")